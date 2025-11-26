
import unittest
import os
import shutil
import subprocess
import zipfile
from lxml import etree
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class TestRedlineStyles(unittest.TestCase):
    def setUp(self):
        self.output_dir = "tests/output"
        self.input_dir = "tests/input"
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.input_dir, exist_ok=True)

    def tearDown(self):
        shutil.rmtree(self.output_dir)
        shutil.rmtree(self.input_dir)

    def create_test_doc(self, filename, content_callback):
        doc = Document()
        content_callback(doc)
        filepath = os.path.join(self.input_dir, filename)
        doc.save(filepath)
        return filepath

    def run_redline(self, before_path, after_path, output_path):
        subprocess.run(
            ["python", "redline_docx/redline_docx_v2.py", before_path, after_path, output_path],
            check=True
        )

    def verify_xml(self, docx_path, verifier_callback):
        with zipfile.ZipFile(docx_path, 'r') as z:
            with z.open('word/document.xml') as f:
                xml_content = f.read()
                root = etree.fromstring(xml_content)
                verifier_callback(root)

    def test_simple_style_preservation(self):
        def before_content(doc):
            doc.add_heading('Heading 1', level=1)
            p = doc.add_paragraph('A plain paragraph with ')
            p.add_run('bold').bold = True
            p.add_run(' and ')
            p.add_run('italic').italic = True
            p.add_run('.')
            doc.add_paragraph('Another paragraph, this one is centered.', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
        def after_content(doc):
            doc.add_heading('Heading 1', level=1)
            p = doc.add_paragraph('A plain paragraph with ')
            p.add_run('bold and styled').bold = True
            p.add_run(' and ')
            p.add_run('italic').italic = True
            p.add_run(' text.')
            doc.add_paragraph('Another paragraph, this one is centered and modified.', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
        before_path = self.create_test_doc("before_simple.docx", before_content)
        after_path = self.create_test_doc("after_simple.docx", after_content)
        output_path = os.path.join(self.output_dir, "output_simple.docx")
        self.run_redline(before_path, after_path, output_path)

        def verify(xml):
            # Check for an insertion
            insertions = xml.xpath("//w:ins", namespaces=xml.nsmap)
            self.assertTrue(len(insertions) > 0, "No insertions found in the output")

        self.verify_xml(output_path, verify)

    def test_list_preservation(self):
        def before_content(doc):
            doc.add_paragraph('First item in list', style='List Bullet')
            doc.add_paragraph('Second item to be removed', style='List Bullet')
            doc.add_paragraph('Third item', style='List Bullet')
        def after_content(doc):
            doc.add_paragraph('First item in list, but modified', style='List Bullet')
            doc.add_paragraph('Third item', style='List Bullet')
        before_path = self.create_test_doc("before_list.docx", before_content)
        after_path = self.create_test_doc("after_list.docx", after_content)
        output_path = os.path.join(self.output_dir, "output_list.docx")
        self.run_redline(before_path, after_path, output_path)

        def verify(xml):
            insertions = xml.xpath("//w:ins", namespaces=xml.nsmap)
            self.assertTrue(len(insertions) > 0, "No insertions found in list test")
            deletions = xml.xpath("//w:del", namespaces=xml.nsmap)
            self.assertTrue(len(deletions) > 0, "No deletions found in list test")

        self.verify_xml(output_path, verify)

    def test_table_preservation(self):
        def before_content(doc):
            table = doc.add_table(rows=2, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Qty'
            hdr_cells[1].text = 'Id'
            hdr_cells[2].text = 'Desc'
            row_cells = table.rows[1].cells
            row_cells[0].text = '1'
            row_cells[1].text = '42'
            row_cells[2].text = 'Item to be deleted'
        def after_content(doc):
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Quantity'
            hdr_cells[1].text = 'Identifier'
            hdr_cells[2].text = 'Description'
        before_path = self.create_test_doc("before_table.docx", before_content)
        after_path = self.create_test_doc("after_table.docx", after_content)
        output_path = os.path.join(self.output_dir, "output_table.docx")
        self.run_redline(before_path, after_path, output_path)

        def verify(xml):
            insertions = xml.xpath("//w:ins", namespaces=xml.nsmap)
            self.assertTrue(len(insertions) > 0, "No insertions found in table test")
            deletions = xml.xpath("//w:del", namespaces=xml.nsmap)
            self.assertTrue(len(deletions) > 0, "No deletions found in table test")

        self.verify_xml(output_path, verify)

if __name__ == '__main__':
    unittest.main()
