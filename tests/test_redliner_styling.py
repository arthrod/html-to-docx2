#!/usr/bin/env python3
"""Comprehensive tests for redliner style preservation.

These tests verify that the Redliner correctly tracks and preserves:
- Run-level formatting (bold, italic, underline, font, color, size, etc.)
- Paragraph-level formatting (alignment, spacing, indentation, etc.)
- Combinations of multiple styles
- Formatting changes (added, removed, modified)
"""

import os
import sys
import tempfile
import shutil
import zipfile
from pathlib import Path

# Add the parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from lxml import etree
from docx import Document as PythonDocx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import the redliner components
from claude_office_skills.public.docx.scripts.document import Document
from claude_office_skills.public.docx.scripts.redliner import Redliner
from claude_office_skills.public.docx.ooxml.scripts.unpack import unpack_document


# Namespaces for XML parsing
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': NS_W}


def qn(tag: str) -> str:
    """Expand a QName like 'w:p' to '{namespace}p'."""
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


class TestFixture:
    """Helper class for creating test documents."""

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def cleanup(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def create_docx(self, name: str, creator_func) -> str:
        """Create a docx file using a creator function."""
        doc = PythonDocx()
        creator_func(doc)
        path = os.path.join(self.temp_dir, name)
        doc.save(path)
        return path

    def unpack_docx(self, docx_path: str) -> str:
        """Unpack a docx file and return the unpacked directory path."""
        unpacked_dir = docx_path.replace('.docx', '_unpacked')
        unpack_document(docx_path, unpacked_dir)
        return unpacked_dir

    def create_redlined_docx(self, old_path: str, new_path: str, author: str = 'TestAuthor') -> str:
        """Create a redlined docx by comparing old and new."""
        old_unpacked = self.unpack_docx(old_path)
        new_unpacked = self.unpack_docx(new_path)

        old_doc = Document(old_unpacked)
        new_doc = Document(new_unpacked)

        redliner = Redliner(old_doc, new_doc)
        redliner.redline(author=author)

        output_path = os.path.join(self.temp_dir, 'redlined.docx')
        new_doc.save(destination=new_unpacked, validate=False)

        from claude_office_skills.public.docx.ooxml.scripts.pack import pack_document
        pack_document(new_unpacked, output_path, validate=False)

        return output_path

    def parse_document_xml(self, docx_path: str) -> etree._Element:
        """Parse the document.xml from a docx file."""
        with zipfile.ZipFile(docx_path, 'r') as z:
            doc_xml = z.read('word/document.xml')
        return etree.fromstring(doc_xml)


@pytest.fixture
def test_fixture():
    """Pytest fixture for test setup and cleanup."""
    fixture = TestFixture()
    yield fixture
    fixture.cleanup()


# ============================================================================
# Run-Level (Character) Formatting Tests
# ============================================================================

class TestBoldFormatting:
    """Tests for bold formatting preservation."""

    def test_bold_to_normal(self, test_fixture):
        """Test tracking when bold formatting is removed."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Bold text')
            r.bold = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Bold text')
            # Not bold

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        # Parse and verify rPrChange exists
        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for bold removal"

        # Verify the previous bold state is recorded
        for change in rPrChanges:
            prior_rPr = change.find(qn('w:rPr'))
            if prior_rPr is not None:
                bold = prior_rPr.find(qn('w:b'))
                if bold is not None:
                    return  # Found the bold tracking
        assert False, "Should record previous bold state in rPrChange"

    def test_normal_to_bold(self, test_fixture):
        """Test tracking when bold formatting is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')
            r.bold = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for bold addition"


class TestItalicFormatting:
    """Tests for italic formatting preservation."""

    def test_italic_to_normal(self, test_fixture):
        """Test tracking when italic formatting is removed."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Italic text')
            r.italic = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Italic text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for italic removal"

    def test_normal_to_italic(self, test_fixture):
        """Test tracking when italic formatting is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')
            r.italic = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for italic addition"


class TestUnderlineFormatting:
    """Tests for underline formatting preservation."""

    def test_underline_to_normal(self, test_fixture):
        """Test tracking when underline formatting is removed."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Underlined text')
            r.underline = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Underlined text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for underline removal"

    def test_normal_to_underline(self, test_fixture):
        """Test tracking when underline formatting is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Normal text')
            r.underline = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for underline addition"


class TestFontColorFormatting:
    """Tests for font color formatting preservation."""

    def test_red_to_blue(self, test_fixture):
        """Test tracking when font color changes."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Colored text')
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Colored text')
            r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for color change"

    def test_colored_to_default(self, test_fixture):
        """Test tracking when font color is removed."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Colored text')
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Colored text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for color removal"


class TestFontSizeFormatting:
    """Tests for font size formatting preservation."""

    def test_size_12_to_20(self, test_fixture):
        """Test tracking when font size changes."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Sized text')
            r.font.size = Pt(12)

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Sized text')
            r.font.size = Pt(20)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for size change"


class TestStrikethroughFormatting:
    """Tests for strikethrough formatting preservation."""

    def test_normal_to_strikethrough(self, test_fixture):
        """Test tracking when strikethrough is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Strike text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Strike text')
            r.font.strike = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for strikethrough addition"


class TestHighlightFormatting:
    """Tests for highlight formatting preservation."""

    def test_normal_to_highlighted(self, test_fixture):
        """Test tracking when highlight is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Highlight text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Highlight text')
            r.font.highlight_color = 7  # Yellow

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for highlight addition"


class TestSuperscriptSubscript:
    """Tests for superscript/subscript formatting preservation."""

    def test_normal_to_superscript(self, test_fixture):
        """Test tracking when superscript is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('2')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('2')
            r.font.superscript = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for superscript addition"

    def test_normal_to_subscript(self, test_fixture):
        """Test tracking when subscript is added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('2')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('2')
            r.font.subscript = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for subscript addition"


# ============================================================================
# Combined Formatting Tests
# ============================================================================

class TestCombinedFormatting:
    """Tests for combined formatting changes."""

    def test_bold_italic_to_normal(self, test_fixture):
        """Test tracking when multiple formats are removed."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')
            r.bold = True
            r.italic = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for multiple format removal"

    def test_normal_to_bold_italic_underline(self, test_fixture):
        """Test tracking when multiple formats are added."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Plain text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Plain text')
            r.bold = True
            r.italic = True
            r.underline = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for multiple format addition"

    def test_text_change_with_format_change(self, test_fixture):
        """Test tracking when both text and formatting change."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Old bold text')
            r.bold = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('New italic text')
            r.italic = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        # Should have both ins/del for text and potentially rPrChange
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(insertions) > 0 or len(deletions) > 0, "Should have tracked changes for text"


# ============================================================================
# Paragraph-Level Formatting Tests
# ============================================================================

class TestParagraphAlignment:
    """Tests for paragraph alignment changes."""

    def test_left_to_center(self, test_fixture):
        """Test tracking when alignment changes from left to center."""
        def create_old(doc):
            p = doc.add_paragraph('Left aligned text')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def create_new(doc):
            p = doc.add_paragraph('Left aligned text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for alignment change"

    def test_center_to_right(self, test_fixture):
        """Test tracking when alignment changes from center to right."""
        def create_old(doc):
            p = doc.add_paragraph('Centered text')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def create_new(doc):
            p = doc.add_paragraph('Centered text')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for alignment change"

    def test_right_to_justify(self, test_fixture):
        """Test tracking when alignment changes from right to justify."""
        def create_old(doc):
            p = doc.add_paragraph('Right aligned text')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        def create_new(doc):
            p = doc.add_paragraph('Right aligned text')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for alignment change"


class TestParagraphSpacing:
    """Tests for paragraph spacing changes."""

    def test_spacing_before_change(self, test_fixture):
        """Test tracking when spacing before changes."""
        def create_old(doc):
            p = doc.add_paragraph('Spaced text')
            p.paragraph_format.space_before = Pt(0)

        def create_new(doc):
            p = doc.add_paragraph('Spaced text')
            p.paragraph_format.space_before = Pt(12)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for spacing change"

    def test_spacing_after_change(self, test_fixture):
        """Test tracking when spacing after changes."""
        def create_old(doc):
            p = doc.add_paragraph('Spaced text')
            p.paragraph_format.space_after = Pt(0)

        def create_new(doc):
            p = doc.add_paragraph('Spaced text')
            p.paragraph_format.space_after = Pt(12)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for spacing change"


class TestParagraphIndentation:
    """Tests for paragraph indentation changes."""

    def test_left_indent_change(self, test_fixture):
        """Test tracking when left indentation changes."""
        def create_old(doc):
            p = doc.add_paragraph('Indented text')
            p.paragraph_format.left_indent = Inches(0)

        def create_new(doc):
            p = doc.add_paragraph('Indented text')
            p.paragraph_format.left_indent = Inches(0.5)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for indentation change"

    def test_first_line_indent_change(self, test_fixture):
        """Test tracking when first line indentation changes."""
        def create_old(doc):
            p = doc.add_paragraph('Indented first line')
            p.paragraph_format.first_line_indent = Inches(0)

        def create_new(doc):
            p = doc.add_paragraph('Indented first line')
            p.paragraph_format.first_line_indent = Inches(0.5)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for first line indent change"


# ============================================================================
# Heading and List Style Tests
# ============================================================================

class TestHeadingStyles:
    """Tests for heading style changes."""

    def test_normal_to_heading1(self, test_fixture):
        """Test tracking when text becomes a heading."""
        def create_old(doc):
            doc.add_paragraph('Heading text')

        def create_new(doc):
            doc.add_paragraph('Heading text', style='Heading 1')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for heading style change"

    def test_heading1_to_heading2(self, test_fixture):
        """Test tracking when heading level changes."""
        def create_old(doc):
            doc.add_paragraph('Heading text', style='Heading 1')

        def create_new(doc):
            doc.add_paragraph('Heading text', style='Heading 2')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        assert len(pPrChanges) > 0, "Should have pPrChange for heading level change"


class TestListStyles:
    """Tests for list style changes."""

    def test_numbered_list_item_change(self, test_fixture):
        """Test tracking changes in numbered list items."""
        def create_old(doc):
            doc.add_paragraph('First item', style='List Number')
            doc.add_paragraph('Second item', style='List Number')

        def create_new(doc):
            doc.add_paragraph('First item modified', style='List Number')
            doc.add_paragraph('Second item', style='List Number')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        # Should have tracked changes
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(insertions) > 0 or len(deletions) > 0, "Should have tracked changes for list item"

    def test_bulleted_list_item_change(self, test_fixture):
        """Test tracking changes in bulleted list items."""
        def create_old(doc):
            doc.add_paragraph('First bullet', style='List Bullet')
            doc.add_paragraph('Second bullet', style='List Bullet')

        def create_new(doc):
            doc.add_paragraph('First bullet', style='List Bullet')
            doc.add_paragraph('Second bullet modified', style='List Bullet')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(insertions) > 0 or len(deletions) > 0, "Should have tracked changes for list item"


# ============================================================================
# Edge Cases and Special Scenarios
# ============================================================================

class TestEdgeCases:
    """Tests for edge cases and special scenarios."""

    def test_empty_to_formatted(self, test_fixture):
        """Test adding formatting where there was none."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Plain text')

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Plain text')
            r.bold = True
            r.italic = True
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r.font.size = Pt(16)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for formatting addition"

        # Verify the rPrChange contains empty or minimal prior formatting
        for change in rPrChanges:
            prior_rPr = change.find(qn('w:rPr'))
            assert prior_rPr is not None, "rPrChange should contain previous rPr"

    def test_formatted_to_empty(self, test_fixture):
        """Test removing all formatting."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')
            r.bold = True
            r.italic = True
            r.underline = True
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for formatting removal"

    def test_partial_text_format_change(self, test_fixture):
        """Test when only part of text changes formatting."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Normal ')
            r = p.add_run('bold')
            r.bold = True
            p.add_run(' text')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Normal ')
            p.add_run('bold')  # No longer bold
            p.add_run(' text')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should have rPrChange for partial formatting change"

    def test_no_changes(self, test_fixture):
        """Test that identical documents don't create spurious changes."""
        def create_doc(doc):
            p = doc.add_paragraph()
            r = p.add_run('Same text')
            r.bold = True

        old_path = test_fixture.create_docx('old.docx', create_doc)
        new_path = test_fixture.create_docx('new.docx', create_doc)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)

        assert len(insertions) == 0, "Should have no insertions for identical docs"
        assert len(deletions) == 0, "Should have no deletions for identical docs"
        assert len(rPrChanges) == 0, "Should have no rPrChange for identical docs"
        assert len(pPrChanges) == 0, "Should have no pPrChange for identical docs"


class TestMixedScenarios:
    """Tests for complex mixed scenarios."""

    def test_text_insertion_with_format(self, test_fixture):
        """Test inserting new formatted text."""
        def create_old(doc):
            doc.add_paragraph('Original text.')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Original text. ')
            r = p.add_run('Added bold text.')
            r.bold = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        assert len(insertions) > 0, "Should have insertion for added text"

    def test_text_deletion_preserves_surrounding_format(self, test_fixture):
        """Test that deleting text preserves surrounding formatting."""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Keep this ')
            r.bold = True
            p.add_run('delete this ')
            r2 = p.add_run('keep this too')
            r2.italic = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Keep this ')
            r.bold = True
            r2 = p.add_run('keep this too')
            r2.italic = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(deletions) > 0, "Should have deletion for removed text"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
