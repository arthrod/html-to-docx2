import os
import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add the Hyperlink style
    style = docx.oxml.shared.OxmlElement('w:rStyle')
    style.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
    rPr.append(style)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    # Manual styling is still useful as a fallback
    r.font.color.rgb = RGBColor(0x05, 0x63, 0xc1)
    r.font.underline = True

    return hyperlink

def create_test_docs():
    """Creates all the test documents."""
    # Create the directory for the test files
    if not os.path.exists('tests/test_files'):
        os.makedirs('tests/test_files')

    # Create the test files
    create_basic_functionality_docs()
    create_character_level_diffs_docs()
    create_table_cell_redlining_docs()
    create_hyperlink_preservation_docs()
    create_style_and_formatting_docs()
    create_style_preservation_docs()
    create_progressive_complexity_docs()


def create_basic_functionality_docs():
    """Creates test documents for basic functionality."""
    # Test Case 1.1: Simple text insertion
    doc_old = Document()
    doc_old.add_paragraph('This is a test.')
    doc_old.save(os.path.join('tests', 'test_files', '1.1_old.docx'))
    doc_new = Document()
    doc_new.add_paragraph('This is a test insertion.')
    doc_new.save('tests/test_files/1.1_new.docx')

    # Test Case 1.2: Simple text deletion
    doc_old = Document()
    doc_old.add_paragraph('This is a test deletion.')
    doc_old.save('tests/test_files/1.2_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test.')
    doc_new.save('tests/test_files/1.2_new.docx')

    # Test Case 1.3: Simple text replacement
    doc_old = Document()
    doc_old.add_paragraph('This is a test replacement.')
    doc_old.save('tests/test_files/1.3_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test substitution.')
    doc_new.save('tests/test_files/1.3_new.docx')

    # Test Case 1.4: No changes
    doc_old = Document()
    doc_old.add_paragraph('This is a test with no changes.')
    doc_old.save('tests/test_files/1.4_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test with no changes.')
    doc_new.save('tests/test_files/1.4_new.docx')


def create_character_level_diffs_docs():
    """Creates test documents for character-level diffs."""
    # Test Case 2.1: Character insertion within a word
    doc_old = Document()
    doc_old.add_paragraph('This is a test caracter.')
    doc_old.save('tests/test_files/2.1_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test character.')
    doc_new.save('tests/test_files/2.1_new.docx')

    # Test Case 2.2: Character deletion within a word
    doc_old = Document()
    doc_old.add_paragraph('This is a test character.')
    doc_old.save('tests/test_files/2.2_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test caracter.')
    doc_new.save('tests/test_files/2.2_new.docx')

    # Test Case 2.3: Character replacement within a word
    doc_old = Document()
    doc_old.add_paragraph('This is a test character.')
    doc_old.save('tests/test_files/2.3_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test charactar.')
    doc_new.save('tests/test_files/2.3_new.docx')


def create_table_cell_redlining_docs():
    """Creates test documents for table cell redlining."""
    # Test Case 3.1: Text change in a table cell
    doc_old = Document()
    table = doc_old.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    doc_old.save('tests/test_files/3.1_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test change.'
    doc_new.save('tests/test_files/3.1_new.docx')

    # Test Case 3.2: Add a row to a table
    doc_old = Document()
    table = doc_old.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    doc_old.save('tests/test_files/3.2_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=2, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    table.cell(1, 0).text = 'This is a new row.'
    doc_new.save('tests/test_files/3.2_new.docx')

    # Test Case 3.3: Delete a row from a table
    doc_old = Document()
    table = doc_old.add_table(rows=2, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    table.cell(1, 0).text = 'This is a row to be deleted.'
    doc_old.save('tests/test_files/3.3_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    doc_new.save('tests/test_files/3.3_new.docx')

    # Test Case 3.4: Add a column to a table
    doc_old = Document()
    table = doc_old.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    doc_old.save('tests/test_files/3.4_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'This is a test.'
    table.cell(0, 1).text = 'This is a new column.'
    doc_new.save('tests/test_files/3.4_new.docx')

    # Test Case 3.5: Delete a column from a table
    doc_old = Document()
    table = doc_old.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'This is a test.'
    table.cell(0, 1).text = 'This is a column to be deleted.'
    doc_old.save('tests/test_files/3.5_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'This is a test.'
    doc_new.save('tests/test_files/3.5_new.docx')

    # Test Case 3.6: Changes in a nested table
    doc_old = Document()
    table = doc_old.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    nested_table = cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = 'This is a test.'
    doc_old.save('tests/test_files/3.6_old.docx')
    doc_new = Document()
    table = doc_new.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    nested_table = cell.add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = 'This is a test change.'
    doc_new.save('tests/test_files/3.6_new.docx')


def create_hyperlink_preservation_docs():
    """Creates test documents for hyperlink preservation."""
    # Test Case 4.1: Unchanged hyperlink
    doc_old = Document()
    p = doc_old.add_paragraph('This is a ')
    add_hyperlink(p, 'hyperlink', 'http://test.com')
    doc_old.save('tests/test_files/4.1_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph('This is a ')
    add_hyperlink(p, 'hyperlink', 'http://test.com')
    doc_new.save('tests/test_files/4.1_new.docx')

    # Test Case 4.2: Change hyperlink text
    doc_old = Document()
    p = doc_old.add_paragraph('This is a ')
    add_hyperlink(p, 'hyperlink', 'http://test.com')
    doc_old.save('tests/test_files/4.2_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph('This is a ')
    add_hyperlink(p, 'changed hyperlink', 'http://test.com')
    doc_new.save('tests/test_files/4.2_new.docx')

    # Test Case 4.3: Add a hyperlink
    doc_old = Document()
    doc_old.add_paragraph('This is a test.')
    doc_old.save('tests/test_files/4.3_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph('This is a ')
    add_hyperlink(p, 'hyperlink', 'http://test.com')
    doc_new.save('tests/test_files/4.3_new.docx')

    # Test Case 4.4: Delete a hyperlink
    doc_old = Document()
    p = doc_old.add_paragraph('This is a ')
    add_hyperlink(p, 'hyperlink', 'http://test.com')
    doc_old.save('tests/test_files/4.4_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('This is a test.')
    doc_new.save('tests/test_files/4.4_new.docx')


def create_style_and_formatting_docs():
    """Creates test documents for style and formatting."""
    # Test Case 5.1: Change bold text
    doc_old = Document()
    p = doc_old.add_paragraph()
    r = p.add_run()
    r.text = 'This is a '
    r = p.add_run()
    r.text = 'bold'
    r.bold = True
    r = p.add_run()
    r.text = ' word.'
    doc_old.save('tests/test_files/5.1_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph()
    r = p.add_run()
    r.text = 'This is a '
    r = p.add_run()
    r.text = 'bold'
    r = p.add_run()
    r.text = ' word.'
    doc_new.save('tests/test_files/5.1_new.docx')

    # Test Case 5.2: Change italic text
    doc_old = Document()
    p = doc_old.add_paragraph()
    r = p.add_run()
    r.text = 'This is an '
    r = p.add_run()
    r.text = 'italic'
    r.italic = True
    r = p.add_run()
    r.text = ' word.'
    doc_old.save('tests/test_files/5.2_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph()
    r = p.add_run()
    r.text = 'This is an '
    r = p.add_run()
    r.text = 'italic'
    r = p.add_run()
    r.text = ' word.'
    doc_new.save('tests/test_files/5.2_new.docx')

    # Test Case 5.3: Change underlined text
    doc_old = Document()
    p = doc_old.add_paragraph()
    r = p.add_run()
    r.text = 'This is an '
    r = p.add_run()
    r.text = 'underlined'
    r.underline = True
    r = p.add_run()
    r.text = ' word.'
    doc_old.save('tests/test_files/5.3_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph()
    r = p.add_run()
    r.text = 'This is an '
    r = p.add_run()
    r.text = 'underlined'
    r = p.add_run()
    r.text = ' word.'
    doc_new.save('tests/test_files/5.3_new.docx')

    # Test Case 5.4: Change font size and color
    doc_old = Document()
    p = doc_old.add_paragraph()
    r = p.add_run()
    r.text = 'This is a word with a '
    r = p.add_run()
    r.text = 'different'
    font = r.font
    font.size = Pt(20)
    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r = p.add_run()
    r.text = ' font.'
    doc_old.save('tests/test_files/5.4_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph()
    r = p.add_run()
    r.text = 'This is a word with a '
    r = p.add_run()
    r.text = 'different'
    r = p.add_run()
    r.text = ' font.'
    doc_new.save('tests/test_files/5.4_new.docx')

    # Test Case 5.5: Change paragraph alignment
    doc_old = Document()
    p = doc_old.add_paragraph('This is a centered paragraph.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_old.save('tests/test_files/5.5_old.docx')
    doc_new = Document()
    p = doc_new.add_paragraph('This is a centered paragraph.')
    doc_new.save('tests/test_files/5.5_new.docx')

    # Test Case 5.6: Changes in a numbered list
    doc_old = Document()
    doc_old.add_paragraph('First item in a numbered list.', style='List Number')
    doc_old.add_paragraph('Second item in a numbered list.', style='List Number')
    doc_old.save('tests/test_files/5.6_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('First item in a numbered list.', style='List Number')
    doc_new.save('tests/test_files/5.6_new.docx')

    # Test Case 5.7: Changes in a bulleted list
    doc_old = Document()
    doc_old.add_paragraph('First item in a bulleted list.', style='List Bullet')
    doc_old.add_paragraph('Second item in a bulleted list.', style='List Bullet')
    doc_old.save('tests/test_files/5.7_old.docx')
    doc_new = Document()
    doc_new.add_paragraph('First item in a bulleted list.', style='List Bullet')
    doc_new.save('tests/test_files/5.7_new.docx')


def create_style_preservation_docs():
    """Creates documents that require merging custom styles during redlining."""
    # Test Case 6.1: Paragraph style only exists in the old document
    doc_old = Document()
    para_style = doc_old.styles.add_style('CustomParaStyle', WD_STYLE_TYPE.PARAGRAPH)
    para_style.font.name = 'Arial'
    para_style.font.size = Pt(14)
    para_style.font.bold = True
    para_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_old.add_paragraph('Styled paragraph with unique style.', style='CustomParaStyle')
    doc_old.save('tests/test_files/6.1_old.docx')

    doc_new = Document()
    doc_new.add_paragraph('Styled paragraph with unique style.')
    doc_new.save('tests/test_files/6.1_new.docx')

    # Test Case 6.2: Character style only exists in the old document
    doc_old = Document()
    char_style = doc_old.styles.add_style('CustomCharStyle', WD_STYLE_TYPE.CHARACTER)
    char_style.font.bold = True
    char_style.font.underline = True
    p = doc_old.add_paragraph()
    run = p.add_run('Styled run with unique character style.')
    run.style = char_style
    doc_old.save('tests/test_files/6.2_old.docx')

    doc_new = Document()
    doc_new.add_paragraph('Styled run with unique character style.')
    doc_new.save('tests/test_files/6.2_new.docx')

    # Test Case 6.3: Second pass uses the previously merged paragraph style
    doc_new = Document()
    doc_new.add_paragraph('Styled paragraph with unique style plus a new sentence.')
    doc_new.save('tests/test_files/6.3_new.docx')

    # Test Case 6.4: Second pass touches the previously merged character style
    doc_new = Document()
    p_new = doc_new.add_paragraph()
    p_new.add_run('Styled run with unique character style after edits. ')
    trailing = p_new.add_run('Extra emphasis segment.')
    trailing.bold = True
    trailing.italic = True
    doc_new.save('tests/test_files/6.4_new.docx')


def create_progressive_complexity_docs():
    """Creates documents that escalate formatting and style complexity."""
    # Test Case 10.1: Multiple formatting runs with different emphasis
    doc_old = Document()
    p_old = doc_old.add_paragraph()
    p_old.add_run('Layered ')
    bold_run_old = p_old.add_run('bold')
    bold_run_old.bold = True
    p_old.add_run(' and ')
    italic_run_old = p_old.add_run('italic')
    italic_run_old.italic = True
    p_old.add_run(' with ')
    underline_old = p_old.add_run('underlines')
    underline_old.underline = True
    p_old.add_run(' in motion.')
    doc_old.save('tests/test_files/10.1_old.docx')

    doc_new = Document()
    p_new = doc_new.add_paragraph()
    p_new.add_run('Layered ')
    bold_run_new = p_new.add_run('bold')
    bold_run_new.bold = True
    combined = p_new.add_run(' emphasis')
    combined.bold = True
    combined.italic = True
    p_new.add_run(' and ')
    italic_run_new = p_new.add_run('italicized')
    italic_run_new.italic = True
    p_new.add_run(' with steady ')
    p_new.add_run('lines. ')
    trailing = p_new.add_run('No underline here now.')
    trailing.underline = False
    doc_new.save('tests/test_files/10.1_new.docx')

    # Test Case 15.1: Custom paragraph and character styles with lists and tables
    doc_old = Document()
    dense_para = doc_old.styles.add_style('DensePara', WD_STYLE_TYPE.PARAGRAPH)
    dense_para.font.name = 'Arial'
    dense_para.font.size = Pt(11)
    dense_para.paragraph_format.space_after = Pt(6)
    accent_char = doc_old.styles.add_style('AccentChar', WD_STYLE_TYPE.CHARACTER)
    accent_char.font.bold = True
    accent_char.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    intro = doc_old.add_paragraph('Dense paragraph with ', style='DensePara')
    intro_run = intro.add_run('accented focus')
    intro_run.style = accent_char
    intro.add_run(' and trailing context.')
    doc_old.add_paragraph('Numbered intro item.', style='List Number')
    doc_old.add_paragraph('Follow-up point.', style='List Number')
    table = doc_old.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Left cell bold'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = 'Right cell colored'
    table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAA, 0x22, 0x22)
    doc_old.save('tests/test_files/15.1_old.docx')

    doc_new = Document()
    doc_new.add_paragraph('Dense paragraph with accented focus and updated tail.')
    doc_new.add_paragraph('Numbered intro item.', style='List Number')
    doc_new.add_paragraph('Additional numbered detail.', style='List Number')
    table_new = doc_new.add_table(rows=1, cols=2)
    table_new.cell(0, 0).text = 'Left cell bold with more detail'
    table_new.cell(0, 0).paragraphs[0].runs[0].bold = True
    table_new.cell(0, 1).text = 'Right cell colored and italicized'
    table_new.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xAA, 0x22, 0x22)
    table_new.cell(0, 1).paragraphs[0].runs[0].italic = True
    doc_new.save('tests/test_files/15.1_new.docx')

    # Test Case 20.1: First pass uses custom styles heavily
    doc_old = Document()
    sequence_para = doc_old.styles.add_style('SequencePara', WD_STYLE_TYPE.PARAGRAPH)
    sequence_para.font.size = Pt(12)
    sequence_para.font.bold = True
    sequence_char = doc_old.styles.add_style('SequenceAccent', WD_STYLE_TYPE.CHARACTER)
    sequence_char.font.underline = True
    sequence_char.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    s_p = doc_old.add_paragraph('Intro with ', style='SequencePara')
    s_run = s_p.add_run('styled accent')
    s_run.style = sequence_char
    s_p.add_run(' and closing note.')
    doc_old.add_paragraph('Secondary line with neutral tone.', style='SequencePara')
    doc_old.save('tests/test_files/20.1_old.docx')

    doc_new = Document()
    doc_new.add_paragraph('Intro with styled accent evolving.', style=None)
    doc_new.add_paragraph('Secondary line with neutral tone and appended clause.')
    doc_new.save('tests/test_files/20.1_new.docx')

    # Test Case 20.2: Second pass builds on the merged style definitions
    doc_new_second = Document()
    doc_new_second.add_paragraph('Intro with styled accent evolving and highlighted sequel.')
    tail_para = doc_new_second.add_paragraph('Secondary line with neutral tone and appended clause plus list:')
    doc_new_second.add_paragraph('First checklist item.', style='List Bullet')
    doc_new_second.add_paragraph('Second checklist item.', style='List Bullet')
    emphasized = tail_para.add_run(' extra accent pull')
    emphasized.bold = True
    emphasized.italic = True
    doc_new_second.save('tests/test_files/20.2_new.docx')

if __name__ == '__main__':
    create_test_docs()
