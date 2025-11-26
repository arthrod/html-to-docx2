#!/usr/bin/env python3
"""
Edge case tests for redline generation with difficulty levels 1-20.

Covers:
- Empty documents and paragraphs
- Text movement
- Multiple runs with varying styles
- Complex formatting combinations
- Whitespace handling
- Special characters
- Large documents
"""

import os
import sys
import shutil
import unittest
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from attempts.attempt_2_fixed import make_redline_docx

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': NS_W}


def qn(tag: str) -> str:
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


class TestEdgeCases(unittest.TestCase):
    """Edge case tests organized by difficulty level 1-20."""

    def setUp(self):
        self.test_dir = os.path.join(os.path.dirname(__file__), 'test_edge_files')
        self.output_dir = os.path.join(os.path.dirname(__file__), 'test_edge_output')
        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)
            os.makedirs(d)

    def tearDown(self):
        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)

    def _get_document_xml(self, docx_path: str) -> etree._Element:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            doc_xml = zf.read('word/document.xml')
        return etree.fromstring(doc_xml)

    def _save_and_redline(self, old_doc, new_doc, name):
        old_path = os.path.join(self.test_dir, f'{name}_old.docx')
        new_path = os.path.join(self.test_dir, f'{name}_new.docx')
        out_path = os.path.join(self.output_dir, f'{name}_out.docx')
        old_doc.save(old_path)
        new_doc.save(new_path)
        make_redline_docx(old_path, new_path, out_path, author='Test')
        return out_path

    # =========================================================================
    # DIFFICULTY 1: Empty documents
    # =========================================================================
    def test_level01_both_empty(self):
        """Level 1: Both documents are empty."""
        old_doc = Document()
        new_doc = Document()
        out_path = self._save_and_redline(old_doc, new_doc, 'level01_both_empty')
        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        # No changes expected
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertEqual(len(ins), 0)
        self.assertEqual(len(dels), 0)

    def test_level01_old_empty_new_has_content(self):
        """Level 1: Old doc empty, new doc has content."""
        old_doc = Document()
        new_doc = Document()
        new_doc.add_paragraph('New content.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level01_old_empty')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        self.assertGreater(len(ins), 0, "Insertion not found when old is empty")

    def test_level01_new_empty_old_has_content(self):
        """Level 1: New doc empty, old doc has content."""
        old_doc = Document()
        old_doc.add_paragraph('Old content.')
        new_doc = Document()
        out_path = self._save_and_redline(old_doc, new_doc, 'level01_new_empty')
        root = self._get_document_xml(out_path)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(dels), 0, "Deletion not found when new is empty")

    # =========================================================================
    # DIFFICULTY 2: Empty paragraphs
    # =========================================================================
    def test_level02_empty_paragraph_added(self):
        """Level 2: Empty paragraph added."""
        old_doc = Document()
        old_doc.add_paragraph('First.')
        new_doc = Document()
        new_doc.add_paragraph('First.')
        new_doc.add_paragraph('')  # Empty paragraph
        new_doc.add_paragraph('Second.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level02_empty_para')
        self.assertTrue(os.path.exists(out_path))

    def test_level02_empty_paragraph_removed(self):
        """Level 2: Empty paragraph removed."""
        old_doc = Document()
        old_doc.add_paragraph('First.')
        old_doc.add_paragraph('')
        old_doc.add_paragraph('Second.')
        new_doc = Document()
        new_doc.add_paragraph('First.')
        new_doc.add_paragraph('Second.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level02_empty_removed')
        self.assertTrue(os.path.exists(out_path))

    # =========================================================================
    # DIFFICULTY 3: Single word changes
    # =========================================================================
    def test_level03_single_word_insert(self):
        """Level 3: Single word inserted."""
        old_doc = Document()
        old_doc.add_paragraph('Hello world.')
        new_doc = Document()
        new_doc.add_paragraph('Hello beautiful world.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level03_word_insert')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)

    def test_level03_single_word_delete(self):
        """Level 3: Single word deleted."""
        old_doc = Document()
        old_doc.add_paragraph('Hello beautiful world.')
        new_doc = Document()
        new_doc.add_paragraph('Hello world.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level03_word_delete')
        root = self._get_document_xml(out_path)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(dels), 0)

    def test_level03_single_word_replace(self):
        """Level 3: Single word replaced."""
        old_doc = Document()
        old_doc.add_paragraph('Hello old world.')
        new_doc = Document()
        new_doc.add_paragraph('Hello new world.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level03_word_replace')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 4: Multiple runs same formatting
    # =========================================================================
    def test_level04_multiple_runs_same_style(self):
        """Level 4: Multiple runs with same formatting - should merge logically."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        p.add_run('Word1 ')
        p.add_run('Word2 ')
        p.add_run('Word3.')
        new_doc = Document()
        new_doc.add_paragraph('Word1 Word2 Word3.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level04_multi_run')
        root = self._get_document_xml(out_path)
        # Should have no changes since text is identical
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertEqual(len(ins), 0)
        self.assertEqual(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 5: Whitespace variations
    # =========================================================================
    def test_level05_leading_whitespace(self):
        """Level 5: Leading whitespace changes."""
        old_doc = Document()
        old_doc.add_paragraph('  Text with leading spaces.')
        new_doc = Document()
        new_doc.add_paragraph('Text with leading spaces.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level05_leading_ws')
        self.assertTrue(os.path.exists(out_path))

    def test_level05_trailing_whitespace(self):
        """Level 5: Trailing whitespace changes."""
        old_doc = Document()
        old_doc.add_paragraph('Text with trailing spaces.  ')
        new_doc = Document()
        new_doc.add_paragraph('Text with trailing spaces.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level05_trailing_ws')
        self.assertTrue(os.path.exists(out_path))

    def test_level05_multiple_spaces(self):
        """Level 5: Multiple spaces between words."""
        old_doc = Document()
        old_doc.add_paragraph('Word1  Word2   Word3.')
        new_doc = Document()
        new_doc.add_paragraph('Word1 Word2 Word3.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level05_multi_space')
        self.assertTrue(os.path.exists(out_path))

    # =========================================================================
    # DIFFICULTY 6: Multiple paragraphs
    # =========================================================================
    def test_level06_paragraph_added_middle(self):
        """Level 6: Paragraph added in the middle."""
        old_doc = Document()
        old_doc.add_paragraph('First paragraph.')
        old_doc.add_paragraph('Third paragraph.')
        new_doc = Document()
        new_doc.add_paragraph('First paragraph.')
        new_doc.add_paragraph('Second paragraph.')
        new_doc.add_paragraph('Third paragraph.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level06_para_middle')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)

    def test_level06_paragraph_deleted_middle(self):
        """Level 6: Paragraph deleted from the middle."""
        old_doc = Document()
        old_doc.add_paragraph('First paragraph.')
        old_doc.add_paragraph('Second paragraph.')
        old_doc.add_paragraph('Third paragraph.')
        new_doc = Document()
        new_doc.add_paragraph('First paragraph.')
        new_doc.add_paragraph('Third paragraph.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level06_para_del')
        root = self._get_document_xml(out_path)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 7: Mixed formatting in same paragraph
    # =========================================================================
    def test_level07_alternating_styles(self):
        """Level 7: Alternating bold/normal in same paragraph."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        p.add_run('Normal ')
        r = p.add_run('bold ')
        r.bold = True
        p.add_run('normal ')
        r = p.add_run('bold.')
        r.bold = True

        new_doc = Document()
        p = new_doc.add_paragraph()
        p.add_run('Normal ')
        p.add_run('bold ')  # No longer bold
        p.add_run('normal ')
        r = p.add_run('bold.')
        r.bold = True

        out_path = self._save_and_redline(old_doc, new_doc, 'level07_alternating')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    # =========================================================================
    # DIFFICULTY 8: Text movement (reordering)
    # =========================================================================
    def test_level08_sentences_swapped(self):
        """Level 8: Two sentences swapped order."""
        old_doc = Document()
        old_doc.add_paragraph('First sentence. Second sentence.')
        new_doc = Document()
        new_doc.add_paragraph('Second sentence. First sentence.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level08_swap')
        root = self._get_document_xml(out_path)
        # Should have both insertions and deletions
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins) + len(dels), 0)

    def test_level08_paragraph_moved(self):
        """Level 8: Paragraph moved to different position."""
        old_doc = Document()
        old_doc.add_paragraph('Para A.')
        old_doc.add_paragraph('Para B.')
        old_doc.add_paragraph('Para C.')
        new_doc = Document()
        new_doc.add_paragraph('Para B.')
        new_doc.add_paragraph('Para A.')
        new_doc.add_paragraph('Para C.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level08_para_move')
        self.assertTrue(os.path.exists(out_path))

    # =========================================================================
    # DIFFICULTY 9: Special characters
    # =========================================================================
    def test_level09_unicode_characters(self):
        """Level 9: Unicode characters (émojis, accents)."""
        old_doc = Document()
        old_doc.add_paragraph('Café résumé naïve.')
        new_doc = Document()
        new_doc.add_paragraph('Cafe resume naive.')
        out_path = self._save_and_redline(old_doc, new_doc, 'level09_unicode')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins) + len(dels), 0)

    def test_level09_special_punctuation(self):
        """Level 9: Special punctuation (quotes, dashes)."""
        old_doc = Document()
        old_doc.add_paragraph('"Hello" - said John.')
        new_doc = Document()
        new_doc.add_paragraph("'Hello' -- said John.")
        out_path = self._save_and_redline(old_doc, new_doc, 'level09_punct')
        self.assertTrue(os.path.exists(out_path))

    # =========================================================================
    # DIFFICULTY 10: Multiple style changes on same text
    # =========================================================================
    def test_level10_style_chain_change(self):
        """Level 10: Text goes from bold to italic to underline."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        r = p.add_run('Styled text.')
        r.bold = True

        new_doc = Document()
        p = new_doc.add_paragraph()
        r = p.add_run('Styled text.')
        r.italic = True
        r.underline = True

        out_path = self._save_and_redline(old_doc, new_doc, 'level10_chain')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    # =========================================================================
    # DIFFICULTY 11: Partial word changes
    # =========================================================================
    def test_level11_character_insert_middle(self):
        """Level 11: Character inserted in middle of word."""
        old_doc = Document()
        old_doc.add_paragraph('recieve')
        new_doc = Document()
        new_doc.add_paragraph('receive')
        out_path = self._save_and_redline(old_doc, new_doc, 'level11_char_mid')
        self.assertTrue(os.path.exists(out_path))

    def test_level11_prefix_suffix_change(self):
        """Level 11: Word prefix or suffix changed."""
        old_doc = Document()
        old_doc.add_paragraph('unhappy unfortunately')
        new_doc = Document()
        new_doc.add_paragraph('happy fortunately')
        out_path = self._save_and_redline(old_doc, new_doc, 'level11_prefix')
        root = self._get_document_xml(out_path)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 12: Run boundary changes
    # =========================================================================
    def test_level12_run_split(self):
        """Level 12: Single run split into multiple with different styles."""
        old_doc = Document()
        old_doc.add_paragraph('One single run of text.')

        new_doc = Document()
        p = new_doc.add_paragraph()
        p.add_run('One ')
        r = p.add_run('single')
        r.bold = True
        p.add_run(' run of ')
        r = p.add_run('text')
        r.italic = True
        p.add_run('.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level12_split')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    def test_level12_run_merge(self):
        """Level 12: Multiple styled runs merged into one plain."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        r = p.add_run('Bold ')
        r.bold = True
        r = p.add_run('italic ')
        r.italic = True
        r = p.add_run('underline.')
        r.underline = True

        new_doc = Document()
        new_doc.add_paragraph('Bold italic underline.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level12_merge')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    # =========================================================================
    # DIFFICULTY 13: Font attribute combinations
    # =========================================================================
    def test_level13_size_and_color_change(self):
        """Level 13: Both font size and color change."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        r = p.add_run('Styled text.')
        r.font.size = Pt(24)
        r.font.color.rgb = RGBColor(255, 0, 0)

        new_doc = Document()
        p = new_doc.add_paragraph()
        r = p.add_run('Styled text.')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 255)

        out_path = self._save_and_redline(old_doc, new_doc, 'level13_size_color')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    # =========================================================================
    # DIFFICULTY 14: Paragraph property combinations
    # =========================================================================
    def test_level14_alignment_and_spacing(self):
        """Level 14: Both alignment and text change."""
        old_doc = Document()
        p = old_doc.add_paragraph('Centered text.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        new_doc = Document()
        p = new_doc.add_paragraph('Left aligned text.')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        out_path = self._save_and_redline(old_doc, new_doc, 'level14_align_text')
        root = self._get_document_xml(out_path)
        ppr_changes = root.xpath('//w:pPrChange', namespaces=NSMAP)
        ins_or_del = root.xpath('//w:ins | //w:del', namespaces=NSMAP)
        self.assertGreater(len(ppr_changes), 0)
        self.assertGreater(len(ins_or_del), 0)

    # =========================================================================
    # DIFFICULTY 15: Many paragraphs with mixed changes
    # =========================================================================
    def test_level15_many_paragraphs_mixed(self):
        """Level 15: 10 paragraphs with various types of changes."""
        old_doc = Document()
        for i in range(10):
            old_doc.add_paragraph(f'Paragraph {i} original content.')

        new_doc = Document()
        new_doc.add_paragraph('Paragraph 0 original content.')  # Same
        new_doc.add_paragraph('Paragraph 1 modified content.')  # Changed
        # Skip paragraph 2 (deleted)
        new_doc.add_paragraph('Paragraph 3 original content.')
        new_doc.add_paragraph('NEW paragraph inserted.')  # Inserted
        new_doc.add_paragraph('Paragraph 4 original content.')
        p = new_doc.add_paragraph()  # Formatting change
        r = p.add_run('Paragraph 5 original content.')
        r.bold = True
        new_doc.add_paragraph('Paragraph 6 original content.')
        new_doc.add_paragraph('Paragraph 7 original content.')
        new_doc.add_paragraph('Paragraph 8 original content.')
        new_doc.add_paragraph('Paragraph 9 original content.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level15_many')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 16: Long paragraphs
    # =========================================================================
    def test_level16_long_paragraph_small_change(self):
        """Level 16: Long paragraph with small change in middle."""
        long_text = 'Word ' * 100
        old_doc = Document()
        old_doc.add_paragraph(long_text + 'original ' + long_text)

        new_doc = Document()
        new_doc.add_paragraph(long_text + 'modified ' + long_text)

        out_path = self._save_and_redline(old_doc, new_doc, 'level16_long')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 17: Interleaved insertions and deletions
    # =========================================================================
    def test_level17_interleaved_changes(self):
        """Level 17: Alternating insertions and deletions."""
        old_doc = Document()
        old_doc.add_paragraph('A B C D E F G H.')

        new_doc = Document()
        new_doc.add_paragraph('A X C Y E Z G W.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level17_interleaved')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 18: Complex run structure with formatting
    # =========================================================================
    def test_level18_complex_formatting(self):
        """Level 18: Complex nested formatting changes."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        r = p.add_run('Bold ')
        r.bold = True
        r = p.add_run('and ')
        r = p.add_run('italic ')
        r.italic = True
        r = p.add_run('mixed ')
        r.bold = True
        r.italic = True
        p.add_run('text.')

        new_doc = Document()
        p = new_doc.add_paragraph()
        p.add_run('Bold ')  # No longer bold
        r = p.add_run('and ')
        r.underline = True  # Added underline
        p.add_run('italic ')  # No longer italic
        r = p.add_run('mixed ')
        r.bold = True  # Still bold, no longer italic
        r = p.add_run('text.')
        r.italic = True  # Added italic

        out_path = self._save_and_redline(old_doc, new_doc, 'level18_complex')
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0)

    # =========================================================================
    # DIFFICULTY 19: Complete document replacement
    # =========================================================================
    def test_level19_complete_replacement(self):
        """Level 19: Completely different documents."""
        old_doc = Document()
        old_doc.add_paragraph('This is the original document.')
        old_doc.add_paragraph('It has multiple paragraphs.')
        old_doc.add_paragraph('All will be replaced.')

        new_doc = Document()
        new_doc.add_paragraph('Completely new content here.')
        new_doc.add_paragraph('Nothing from the original remains.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level19_replace')
        root = self._get_document_xml(out_path)
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        self.assertGreater(len(ins), 0)
        self.assertGreater(len(dels), 0)

    # =========================================================================
    # DIFFICULTY 20: Stress test - many changes
    # =========================================================================
    def test_level20_stress_test(self):
        """Level 20: Large document with many changes of all types."""
        old_doc = Document()
        for i in range(20):
            p = old_doc.add_paragraph()
            if i % 3 == 0:
                r = p.add_run(f'Paragraph {i} with bold.')
                r.bold = True
            elif i % 3 == 1:
                r = p.add_run(f'Paragraph {i} with italic.')
                r.italic = True
            else:
                p.add_run(f'Paragraph {i} plain text.')

        new_doc = Document()
        for i in range(25):
            if i == 5:
                continue  # Delete paragraph 5
            p = new_doc.add_paragraph()
            if i < 20:
                if i % 4 == 0:
                    # Change formatting
                    r = p.add_run(f'Paragraph {i} with italic.')
                    r.italic = True
                elif i % 4 == 1:
                    # Keep same
                    r = p.add_run(f'Paragraph {i} with italic.')
                    r.italic = True
                elif i % 4 == 2:
                    # Modify text
                    p.add_run(f'Paragraph {i} modified text.')
                else:
                    # Keep same
                    p.add_run(f'Paragraph {i} plain text.')
            else:
                # New paragraphs
                p.add_run(f'New paragraph {i}.')

        out_path = self._save_and_redline(old_doc, new_doc, 'level20_stress')
        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        # Should have various types of changes
        ins = root.xpath('//w:ins', namespaces=NSMAP)
        dels = root.xpath('//w:del', namespaces=NSMAP)
        rpr = root.xpath('//w:rPrChange', namespaces=NSMAP)
        total_changes = len(ins) + len(dels) + len(rpr)
        self.assertGreater(total_changes, 5, "Expected many changes in stress test")


if __name__ == '__main__':
    unittest.main()
