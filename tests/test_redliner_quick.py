#!/usr/bin/env python3
"""Quick verification tests for redliner style preservation."""

import os
import sys
import tempfile
import shutil
import zipfile
from pathlib import Path

# Add the parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from lxml import etree
from docx import Document as PythonDocx
from docx.shared import Pt, RGBColor

from claude_office_skills.public.docx.scripts.document import Document
from claude_office_skills.public.docx.scripts.redliner import Redliner
from claude_office_skills.public.docx.ooxml.scripts.unpack import unpack_document
from claude_office_skills.public.docx.ooxml.scripts.pack import pack_document


NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': NS_W}


def run_test(test_name, old_creator, new_creator, expected_check):
    """Run a single test case."""
    temp_dir = tempfile.mkdtemp()
    try:
        # Create old document
        old_doc = PythonDocx()
        old_creator(old_doc)
        old_path = os.path.join(temp_dir, 'old.docx')
        old_doc.save(old_path)

        # Create new document
        new_doc = PythonDocx()
        new_creator(new_doc)
        new_path = os.path.join(temp_dir, 'new.docx')
        new_doc.save(new_path)

        # Unpack documents
        old_unpacked = os.path.join(temp_dir, 'old_unpacked')
        new_unpacked = os.path.join(temp_dir, 'new_unpacked')
        unpack_document(old_path, old_unpacked)
        unpack_document(new_path, new_unpacked)

        # Create Document objects and run redliner
        old_doc_obj = Document(old_unpacked)
        new_doc_obj = Document(new_unpacked)

        redliner = Redliner(old_doc_obj, new_doc_obj)
        redliner.redline(author='TestAuthor')

        # Save the result
        new_doc_obj.save(destination=new_unpacked, validate=False)
        redlined_path = os.path.join(temp_dir, 'redlined.docx')
        pack_document(new_unpacked, redlined_path, validate=False)

        # Parse and check
        with zipfile.ZipFile(redlined_path, 'r') as z:
            doc_xml = z.read('word/document.xml')
        doc = etree.fromstring(doc_xml)

        result = expected_check(doc)
        if result:
            print(f"  PASS: {test_name}")
            return True
        else:
            print(f"  FAIL: {test_name}")
            return False

    except Exception as e:
        print(f"  ERROR: {test_name} - {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def main():
    """Run all quick verification tests."""
    print("Running Redliner Style Preservation Tests")
    print("=" * 50)

    passed = 0
    failed = 0

    # Test 1: Bold to Normal - should have rPrChange
    def old_bold(doc):
        p = doc.add_paragraph()
        r = p.add_run('Bold text')
        r.bold = True

    def new_normal(doc):
        p = doc.add_paragraph()
        r = p.add_run('Bold text')

    def check_rpr_change(doc):
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        return len(rPrChanges) > 0

    if run_test("Bold to Normal formatting change", old_bold, new_normal, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 2: Normal to Bold - should have rPrChange
    def old_normal(doc):
        p = doc.add_paragraph()
        r = p.add_run('Normal text')

    def new_bold(doc):
        p = doc.add_paragraph()
        r = p.add_run('Normal text')
        r.bold = True

    if run_test("Normal to Bold formatting change", old_normal, new_bold, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 3: Italic to Normal
    def old_italic(doc):
        p = doc.add_paragraph()
        r = p.add_run('Italic text')
        r.italic = True

    if run_test("Italic to Normal formatting change", old_italic, new_normal, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 4: Normal to Italic
    def new_italic(doc):
        p = doc.add_paragraph()
        r = p.add_run('Normal text')
        r.italic = True

    if run_test("Normal to Italic formatting change", old_normal, new_italic, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 5: Font color change
    def old_red(doc):
        p = doc.add_paragraph()
        r = p.add_run('Colored text')
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    def new_blue(doc):
        p = doc.add_paragraph()
        r = p.add_run('Colored text')
        r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    if run_test("Red to Blue font color change", old_red, new_blue, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 6: Font size change
    def old_size12(doc):
        p = doc.add_paragraph()
        r = p.add_run('Sized text')
        r.font.size = Pt(12)

    def new_size20(doc):
        p = doc.add_paragraph()
        r = p.add_run('Sized text')
        r.font.size = Pt(20)

    if run_test("Font size 12pt to 20pt change", old_size12, new_size20, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 7: Multiple formats added
    def old_plain(doc):
        p = doc.add_paragraph()
        r = p.add_run('Plain text')

    def new_formatted(doc):
        p = doc.add_paragraph()
        r = p.add_run('Plain text')
        r.bold = True
        r.italic = True
        r.underline = True

    if run_test("Plain to Bold+Italic+Underline", old_plain, new_formatted, check_rpr_change):
        passed += 1
    else:
        failed += 1

    # Test 8: Text insertion (should have w:ins)
    def old_text(doc):
        doc.add_paragraph('Original text.')

    def new_text_added(doc):
        doc.add_paragraph('Original text. Added new text.')

    def check_insertion(doc):
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        return len(insertions) > 0

    if run_test("Text insertion tracking", old_text, new_text_added, check_insertion):
        passed += 1
    else:
        failed += 1

    # Test 9: Text deletion (should have w:del)
    def old_text_long(doc):
        doc.add_paragraph('Original text with some extra words.')

    def new_text_short(doc):
        doc.add_paragraph('Original text.')

    def check_deletion(doc):
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        return len(deletions) > 0

    if run_test("Text deletion tracking", old_text_long, new_text_short, check_deletion):
        passed += 1
    else:
        failed += 1

    # Test 10: Paragraph alignment change (should have pPrChange)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def old_left(doc):
        p = doc.add_paragraph('Left aligned text')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def new_center(doc):
        p = doc.add_paragraph('Left aligned text')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def check_ppr_change(doc):
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        return len(pPrChanges) > 0

    if run_test("Paragraph alignment change (left to center)", old_left, new_center, check_ppr_change):
        passed += 1
    else:
        failed += 1

    # Test 11: No changes - should have no tracked changes
    def same_text(doc):
        p = doc.add_paragraph()
        r = p.add_run('Same text')
        r.bold = True

    def check_no_changes(doc):
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        pPrChanges = doc.xpath('.//w:pPrChange', namespaces=NSMAP)
        return len(insertions) == 0 and len(deletions) == 0 and len(rPrChanges) == 0 and len(pPrChanges) == 0

    if run_test("No changes - identical documents", same_text, same_text, check_no_changes):
        passed += 1
    else:
        failed += 1

    # Test 12: Text change with formatting change
    def old_bold_text(doc):
        p = doc.add_paragraph()
        r = p.add_run('Old bold text')
        r.bold = True

    def new_italic_text(doc):
        p = doc.add_paragraph()
        r = p.add_run('New italic text')
        r.italic = True

    def check_text_and_format_change(doc):
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        return len(insertions) > 0 or len(deletions) > 0

    if run_test("Text change with formatting change", old_bold_text, new_italic_text, check_text_and_format_change):
        passed += 1
    else:
        failed += 1

    print("=" * 50)
    print(f"Results: {passed} passed, {failed} failed")
    return failed == 0


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
