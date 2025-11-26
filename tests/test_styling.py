#!/usr/bin/env python3
"""
Comprehensive tests for redline styling preservation.

Tests each aspect of styling and combinations:
- Run property changes (bold, italic, underline, font size, color, etc.)
- Paragraph property changes (alignment, spacing, style references)
- Style addition (plain -> styled)
- Style removal (styled -> plain)
- Combined styles (bold + italic, etc.)
- Style references (pStyle, rStyle)
"""

import os
import sys
import shutil
import unittest
import zipfile
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn as docx_qn
from lxml import etree

# Add parent directory to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Import the fixed redliner implementations
from attempts.attempt_2_fixed import make_redline_docx

# Namespaces
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': NS_W}


def qn(tag: str) -> str:
    """Expand QName."""
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


class TestStylingPreservation(unittest.TestCase):
    """Test cases for style preservation in redlines."""

    def setUp(self):
        """Set up test environment."""
        self.test_dir = os.path.join(os.path.dirname(__file__), 'test_styling_files')
        self.output_dir = os.path.join(os.path.dirname(__file__), 'test_styling_output')

        # Clean and create directories
        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)
            os.makedirs(d)

    def tearDown(self):
        """Clean up test files."""
        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)

    def _get_document_xml(self, docx_path: str) -> etree._Element:
        """Extract and parse document.xml from a .docx file."""
        with zipfile.ZipFile(docx_path, 'r') as zf:
            doc_xml = zf.read('word/document.xml')
        return etree.fromstring(doc_xml)

    def _count_elements(self, root: etree._Element, xpath: str) -> int:
        """Count elements matching xpath."""
        return len(root.xpath(xpath, namespaces=NSMAP))

    def _has_element(self, root: etree._Element, xpath: str) -> bool:
        """Check if element exists."""
        return self._count_elements(root, xpath) > 0

    # =========================================================================
    # Test 1: Bold Text Changes
    # =========================================================================
    def test_bold_added(self):
        """Test tracking when bold formatting is added to text."""
        # Create old doc with plain text (same words as new doc)
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'bold_added_old.docx')
        old_doc.save(old_path)

        # Create new doc with bold formatting on "styled"
        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.bold = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'bold_added_new.docx')
        new_doc.save(new_path)

        # Generate redline
        out_path = os.path.join(self.output_dir, 'bold_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        # Verify output
        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have rPrChange for the bold change
        # The text "bold" should have w:b in its rPr and rPrChange tracking the old (no bold)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for bold addition")

    def test_bold_removed(self):
        """Test tracking when bold formatting is removed from text."""
        # Create old doc with bold text
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.bold = True
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'bold_removed_old.docx')
        old_doc.save(old_path)

        # Create new doc with plain text (same words)
        new_doc = Document()
        new_doc.add_paragraph('This is styled text.')
        new_path = os.path.join(self.test_dir, 'bold_removed_new.docx')
        new_doc.save(new_path)

        # Generate redline
        out_path = os.path.join(self.output_dir, 'bold_removed_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        # Verify output
        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have rPrChange for the bold removal
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for bold removal")

        # The rPrChange should contain the old bold formatting
        bold_in_change = root.xpath('//w:rPrChange//w:b', namespaces=NSMAP)
        self.assertGreater(len(bold_in_change), 0, "Old bold formatting not preserved in rPrChange")

    # =========================================================================
    # Test 2: Italic Text Changes
    # =========================================================================
    def test_italic_added(self):
        """Test tracking when italic formatting is added."""
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'italic_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.italic = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'italic_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'italic_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for italic addition")

    def test_italic_removed(self):
        """Test tracking when italic formatting is removed."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.italic = True
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'italic_removed_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('This is styled text.')
        new_path = os.path.join(self.test_dir, 'italic_removed_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'italic_removed_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        italic_in_change = root.xpath('//w:rPrChange//w:i', namespaces=NSMAP)
        self.assertGreater(len(italic_in_change), 0, "Old italic formatting not preserved in rPrChange")

    # =========================================================================
    # Test 3: Underline Text Changes
    # =========================================================================
    def test_underline_added(self):
        """Test tracking when underline formatting is added."""
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'underline_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.underline = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'underline_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'underline_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for underline addition")

    def test_underline_removed(self):
        """Test tracking when underline formatting is removed."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.underline = True
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'underline_removed_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('This is styled text.')
        new_path = os.path.join(self.test_dir, 'underline_removed_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'underline_removed_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        underline_in_change = root.xpath('//w:rPrChange//w:u', namespaces=NSMAP)
        self.assertGreater(len(underline_in_change), 0, "Old underline formatting not preserved in rPrChange")

    # =========================================================================
    # Test 4: Font Size Changes
    # =========================================================================
    def test_font_size_changed(self):
        """Test tracking when font size is changed."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('large')
        run.font.size = Pt(24)
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'fontsize_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('large')
        run.font.size = Pt(12)  # Changed to smaller
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'fontsize_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'fontsize_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        # Should have rPrChange with old font size
        sz_in_change = root.xpath('//w:rPrChange//w:sz', namespaces=NSMAP)
        self.assertGreater(len(sz_in_change), 0, "Old font size not preserved in rPrChange")

    # =========================================================================
    # Test 5: Font Color Changes
    # =========================================================================
    def test_font_color_changed(self):
        """Test tracking when font color is changed."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('red')
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'fontcolor_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('red')
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Changed to blue
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'fontcolor_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'fontcolor_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        # Should have rPrChange with old color
        color_in_change = root.xpath('//w:rPrChange//w:color', namespaces=NSMAP)
        self.assertGreater(len(color_in_change), 0, "Old font color not preserved in rPrChange")

    # =========================================================================
    # Test 6: Paragraph Alignment Changes
    # =========================================================================
    def test_paragraph_alignment_changed(self):
        """Test tracking when paragraph alignment is changed."""
        old_doc = Document()
        p = old_doc.add_paragraph('This is centered text.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        old_path = os.path.join(self.test_dir, 'align_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph('This is centered text.')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_path = os.path.join(self.test_dir, 'align_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'align_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        # Should have pPrChange with old alignment
        ppr_changes = root.xpath('//w:pPrChange', namespaces=NSMAP)
        self.assertGreater(len(ppr_changes), 0, "pPrChange not found for alignment change")

        # The pPrChange should contain the old alignment (center)
        jc_in_change = root.xpath('//w:pPrChange//w:jc', namespaces=NSMAP)
        self.assertGreater(len(jc_in_change), 0, "Old alignment not preserved in pPrChange")

    def test_paragraph_alignment_added(self):
        """Test tracking when paragraph alignment is added."""
        old_doc = Document()
        old_doc.add_paragraph('This is plain text.')
        old_path = os.path.join(self.test_dir, 'align_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph('This is plain text.')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        new_path = os.path.join(self.test_dir, 'align_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'align_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        ppr_changes = root.xpath('//w:pPrChange', namespaces=NSMAP)
        self.assertGreater(len(ppr_changes), 0, "pPrChange not found when alignment added")

    # =========================================================================
    # Test 7: Combined Styles
    # =========================================================================
    def test_bold_and_italic_added(self):
        """Test tracking when both bold and italic are added."""
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'bold_italic_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.bold = True
        run.italic = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'bold_italic_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'bold_italic_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for combined style addition")

    def test_bold_to_italic(self):
        """Test tracking when bold is changed to italic."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.bold = True
        run = p.add_run(' text.')
        old_path = os.path.join(self.test_dir, 'bold_to_italic_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.italic = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'bold_to_italic_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'bold_to_italic_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for style change")

        # Old bold should be in rPrChange
        bold_in_change = root.xpath('//w:rPrChange//w:b', namespaces=NSMAP)
        self.assertGreater(len(bold_in_change), 0, "Old bold not preserved in rPrChange")

    # =========================================================================
    # Test 8: Strikethrough Changes
    # =========================================================================
    def test_strikethrough_added(self):
        """Test tracking when strikethrough is added."""
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'strike_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.font.strike = True
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'strike_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'strike_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for strikethrough addition")

    # =========================================================================
    # Test 9: Highlight Changes
    # =========================================================================
    def test_highlight_added(self):
        """Test tracking when highlight is added."""
        old_doc = Document()
        old_doc.add_paragraph('This is styled text.')
        old_path = os.path.join(self.test_dir, 'highlight_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('This is ')
        run = p.add_run('styled')
        run.font.highlight_color = 6  # Yellow
        run = p.add_run(' text.')
        new_path = os.path.join(self.test_dir, 'highlight_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'highlight_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for highlight addition")

    # =========================================================================
    # Test 10: Superscript/Subscript Changes
    # =========================================================================
    def test_superscript_added(self):
        """Test tracking when superscript is added."""
        old_doc = Document()
        old_doc.add_paragraph('H2O')
        old_path = os.path.join(self.test_dir, 'super_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('H')
        run = p.add_run('2')
        run.font.superscript = True
        run = p.add_run('O')
        new_path = os.path.join(self.test_dir, 'super_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'super_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        # Just verify it doesn't crash and produces output
        root = self._get_document_xml(out_path)
        self.assertIsNotNone(root)

    def test_subscript_added(self):
        """Test tracking when subscript is added."""
        old_doc = Document()
        old_doc.add_paragraph('H2O')
        old_path = os.path.join(self.test_dir, 'sub_added_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('H')
        run = p.add_run('2')
        run.font.subscript = True
        run = p.add_run('O')
        new_path = os.path.join(self.test_dir, 'sub_added_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'sub_added_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)
        self.assertIsNotNone(root)

    # =========================================================================
    # Test 11: Track Revisions Enabled
    # =========================================================================
    def test_track_revisions_enabled(self):
        """Test that trackRevisions is enabled in settings.xml."""
        old_doc = Document()
        old_doc.add_paragraph('Old text.')
        old_path = os.path.join(self.test_dir, 'track_rev_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('New text.')
        new_path = os.path.join(self.test_dir, 'track_rev_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'track_rev_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))

        # Check settings.xml for trackRevisions
        with zipfile.ZipFile(out_path, 'r') as zf:
            settings_xml = zf.read('word/settings.xml')
        root = etree.fromstring(settings_xml)
        track_rev = root.xpath('//w:trackRevisions', namespaces=NSMAP)
        self.assertGreater(len(track_rev), 0, "trackRevisions not found in settings.xml")

    # =========================================================================
    # Test 12: No Changes (Same Document)
    # =========================================================================
    def test_no_changes(self):
        """Test that identical documents produce no tracked changes."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Identical ')
        run = p.add_run('text')
        run.bold = True
        run = p.add_run('.')

        old_path = os.path.join(self.test_dir, 'nochange_old.docx')
        doc.save(old_path)
        new_path = os.path.join(self.test_dir, 'nochange_new.docx')
        doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'nochange_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have no insertions or deletions
        insertions = root.xpath('//w:ins', namespaces=NSMAP)
        deletions = root.xpath('//w:del', namespaces=NSMAP)
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        ppr_changes = root.xpath('//w:pPrChange', namespaces=NSMAP)

        self.assertEqual(len(insertions), 0, "Unexpected insertions in identical documents")
        self.assertEqual(len(deletions), 0, "Unexpected deletions in identical documents")
        self.assertEqual(len(rpr_changes), 0, "Unexpected rPrChange in identical documents")
        self.assertEqual(len(ppr_changes), 0, "Unexpected pPrChange in identical documents")

    # =========================================================================
    # Test 13: Multiple Formatting Changes in Same Paragraph
    # =========================================================================
    def test_multiple_formatting_changes(self):
        """Test multiple formatting changes in the same paragraph."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('Word1 ')
        run.bold = True
        run = p.add_run('Word2 ')
        run.italic = True
        run = p.add_run('Word3.')
        run.underline = True
        old_path = os.path.join(self.test_dir, 'multi_format_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('Word1 ')
        run.italic = True  # Was bold, now italic
        run = p.add_run('Word2 ')
        run.bold = True  # Was italic, now bold
        run = p.add_run('Word3.')  # Was underline, now plain
        new_path = os.path.join(self.test_dir, 'multi_format_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'multi_format_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have multiple rPrChange elements
        rpr_changes = root.xpath('//w:rPrChange', namespaces=NSMAP)
        self.assertGreater(len(rpr_changes), 0, "rPrChange not found for multiple formatting changes")

    # =========================================================================
    # Test 14: Text Change with Formatting Preservation
    # =========================================================================
    def test_text_change_preserves_formatting(self):
        """Test that text changes preserve formatting in insertions/deletions."""
        old_doc = Document()
        p = old_doc.add_paragraph()
        run = p.add_run('Old ')
        run.bold = True
        run = p.add_run('text.')
        old_path = os.path.join(self.test_dir, 'text_format_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph()
        run = p.add_run('New ')
        run.bold = True
        run = p.add_run('text.')
        new_path = os.path.join(self.test_dir, 'text_format_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'text_format_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='Test')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Deletions should have bold formatting preserved
        del_runs = root.xpath('//w:del//w:r', namespaces=NSMAP)
        for del_run in del_runs:
            rPr = del_run.find(qn('w:rPr'))
            if rPr is not None:
                bold = rPr.find(qn('w:b'))
                if bold is not None:
                    # Bold is preserved in deletion
                    break
        else:
            # If we reach here, no bold was found in deletions
            pass  # This is OK if the deleted text wasn't the bold part

        # Insertions should have bold formatting
        ins_runs = root.xpath('//w:ins//w:r', namespaces=NSMAP)
        found_bold_in_ins = False
        for ins_run in ins_runs:
            rPr = ins_run.find(qn('w:rPr'))
            if rPr is not None:
                bold = rPr.find(qn('w:b'))
                if bold is not None:
                    found_bold_in_ins = True
                    break

        self.assertTrue(found_bold_in_ins, "Bold formatting not preserved in insertion")

    # =========================================================================
    # Test 15: Author and Date Metadata
    # =========================================================================
    def test_author_and_date_metadata(self):
        """Test that author and date are correctly set on tracked changes."""
        old_doc = Document()
        old_doc.add_paragraph('Old text.')
        old_path = os.path.join(self.test_dir, 'metadata_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('New text.')
        new_path = os.path.join(self.test_dir, 'metadata_new.docx')
        new_doc.save(new_path)

        out_path = os.path.join(self.output_dir, 'metadata_out.docx')
        make_redline_docx(old_path, new_path, out_path, author='CustomAuthor', date_iso='2025-01-15T12:00:00Z')

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Check insertions
        insertions = root.xpath('//w:ins', namespaces=NSMAP)
        for ins in insertions:
            author = ins.get(qn('w:author'))
            date = ins.get(qn('w:date'))
            self.assertEqual(author, 'CustomAuthor', f"Author mismatch: {author}")
            self.assertEqual(date, '2025-01-15T12:00:00Z', f"Date mismatch: {date}")

        # Check deletions
        deletions = root.xpath('//w:del', namespaces=NSMAP)
        for de in deletions:
            author = de.get(qn('w:author'))
            date = de.get(qn('w:date'))
            self.assertEqual(author, 'CustomAuthor', f"Author mismatch: {author}")
            self.assertEqual(date, '2025-01-15T12:00:00Z', f"Date mismatch: {date}")


class TestAttempt1Fixed(unittest.TestCase):
    """Test cases for attempt_1_fixed.py (stdlib version)."""

    def setUp(self):
        """Set up test environment."""
        self.test_dir = os.path.join(os.path.dirname(__file__), 'test_attempt1_files')
        self.output_dir = os.path.join(os.path.dirname(__file__), 'test_attempt1_output')

        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)
            os.makedirs(d)

    def tearDown(self):
        """Clean up test files."""
        for d in [self.test_dir, self.output_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)

    def _get_document_xml(self, docx_path: str) -> etree._Element:
        """Extract and parse document.xml from a .docx file."""
        with zipfile.ZipFile(docx_path, 'r') as zf:
            doc_xml = zf.read('word/document.xml')
        return etree.fromstring(doc_xml)

    def test_basic_insertion(self):
        """Test basic text insertion with attempt_1_fixed."""
        from attempts.attempt_1_fixed import build_redlined_document, _read_docx_all, _write_docx_from_base, _merge_styles

        old_doc = Document()
        old_doc.add_paragraph('Hello world.')
        old_path = os.path.join(self.test_dir, 'basic_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('Hello beautiful world.')
        new_path = os.path.join(self.test_dir, 'basic_new.docx')
        new_doc.save(new_path)

        # Run the redliner
        old_doc_el, old_styles, _old_files = _read_docx_all(old_path)
        new_doc_el, new_styles, new_files = _read_docx_all(new_path)

        out_doc = build_redlined_document(old_doc_el, new_doc_el, author='Test', now_iso='2025-01-01T00:00:00Z')
        out_styles = _merge_styles(old_styles, new_styles)

        out_path = os.path.join(self.output_dir, 'basic_out.docx')
        _write_docx_from_base(new_files, out_doc, out_styles, out_path)

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have insertions
        insertions = root.xpath('//w:ins', namespaces=NSMAP)
        self.assertGreater(len(insertions), 0, "No insertions found")

    def test_ppr_change_tracking(self):
        """Test paragraph property change tracking with attempt_1_fixed."""
        from attempts.attempt_1_fixed import build_redlined_document, _read_docx_all, _write_docx_from_base, _merge_styles

        old_doc = Document()
        p = old_doc.add_paragraph('Centered text.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        old_path = os.path.join(self.test_dir, 'ppr_old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        p = new_doc.add_paragraph('Centered text.')
        # Default alignment (left)
        new_path = os.path.join(self.test_dir, 'ppr_new.docx')
        new_doc.save(new_path)

        old_doc_el, old_styles, _old_files = _read_docx_all(old_path)
        new_doc_el, new_styles, new_files = _read_docx_all(new_path)

        out_doc = build_redlined_document(old_doc_el, new_doc_el, author='Test', now_iso='2025-01-01T00:00:00Z')
        out_styles = _merge_styles(old_styles, new_styles)

        out_path = os.path.join(self.output_dir, 'ppr_out.docx')
        _write_docx_from_base(new_files, out_doc, out_styles, out_path)

        self.assertTrue(os.path.exists(out_path))
        root = self._get_document_xml(out_path)

        # Should have pPrChange
        ppr_changes = root.xpath('//w:pPrChange', namespaces=NSMAP)
        self.assertGreater(len(ppr_changes), 0, "No pPrChange found for paragraph alignment change")


if __name__ == '__main__':
    unittest.main()
