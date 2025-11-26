"""Style preservation tests for redline generation.

These tests focus on ensuring paragraph/run styles and formatting combinations are
retained in the generated redline output, including scenarios where the old
document contains custom styles that are absent from the new document.
"""

from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path
import warnings

from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from lxml import etree

from redline_docx.redline_docx_enhanced import make_redline_docx


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# python-docx leaves temporary unpacked files open; suppress resource warnings to keep test output clean
warnings.filterwarnings("ignore", category=ResourceWarning)


def _read_part(docx_path: Path, part: str) -> etree._Element:
    with zipfile.ZipFile(docx_path) as zf:
        return etree.fromstring(zf.read(part))


class RedlineStylePreservationTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.tmpdir = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_custom_styles_merged_when_only_in_old_document(self) -> None:
        """Custom paragraph styles from the old document should be merged into the output."""

        old_doc = DocxDocument()
        custom_style = old_doc.styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)
        custom_style.font.name = "Arial"
        custom_style.font.size = Pt(18)
        old_doc.add_paragraph("Styled section", style="CustomHeading")
        old_path = self.tmpdir / "old.docx"
        old_doc.save(old_path)

        new_doc = DocxDocument()
        new_doc.add_paragraph("Styled section")
        new_path = self.tmpdir / "new.docx"
        new_doc.save(new_path)

        out_path = self.tmpdir / "out.docx"
        make_redline_docx(str(old_path), str(new_path), str(out_path), author="Tester", date_iso="2025-01-01T00:00:00Z")

        styles_root = _read_part(out_path, "word/styles.xml")
        self.assertTrue(
            styles_root.xpath("//w:style[@w:styleId='CustomHeading']", namespaces=NS),
            "Custom style from the old document should be copied into the output styles.xml",
        )

        doc_root = _read_part(out_path, "word/document.xml")
        ppr_changes = doc_root.xpath("//w:pPrChange[w:pPr/w:pStyle[@w:val='CustomHeading']]", namespaces=NS)
        self.assertTrue(ppr_changes, "Paragraph change should reference the old custom style for tracking.")

    def test_run_formatting_changes_capture_old_and_new_styles(self) -> None:
        """Formatting-only changes should keep new styling and record prior formatting in rPrChange."""

        old_doc = DocxDocument()
        old_run = old_doc.add_paragraph().add_run("FormatCheck")
        old_run.font.bold = True
        old_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        old_run.font.underline = True
        old_path = self.tmpdir / "old_format.docx"
        old_doc.save(old_path)

        new_doc = DocxDocument()
        new_run = new_doc.add_paragraph().add_run("FormatCheck")
        new_run.font.italic = True
        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        new_run.font.size = Pt(20)
        new_run.font.name = "Georgia"
        new_path = self.tmpdir / "new_format.docx"
        new_doc.save(new_path)

        out_path = self.tmpdir / "out_format.docx"
        make_redline_docx(str(old_path), str(new_path), str(out_path), author="Tester", date_iso="2025-01-02T00:00:00Z")

        doc_root = _read_part(out_path, "word/document.xml")
        run = doc_root.xpath("//w:p[1]/w:r[1]", namespaces=NS)[0]
        rpr = run.find("w:rPr", namespaces=NS)
        self.assertIsNotNone(rpr, "New run formatting should be present")
        self.assertIsNotNone(rpr.find("w:i", namespaces=NS))
        self.assertIsNotNone(rpr.find("w:highlight", namespaces=NS))
        size_el = rpr.find("w:sz", namespaces=NS)
        self.assertEqual(size_el.get(f"{{{NS['w']}}}val"), "40")
        fonts = rpr.find("w:rFonts", namespaces=NS)
        self.assertEqual(fonts.get(f"{{{NS['w']}}}ascii"), "Georgia")

        rpr_change = rpr.find("w:rPrChange", namespaces=NS)
        self.assertIsNotNone(rpr_change, "Formatting diff should emit rPrChange with prior style")
        old_rpr = rpr_change.find("w:rPr", namespaces=NS)
        self.assertIsNotNone(old_rpr.find("w:b", namespaces=NS))
        self.assertIsNotNone(old_rpr.find("w:u", namespaces=NS))
        old_color = old_rpr.find("w:color", namespaces=NS)
        self.assertEqual(old_color.get(f"{{{NS['w']}}}val"), "CC0000")

    def test_inserted_runs_keep_combined_formatting(self) -> None:
        """Inserted text should retain all formatting attributes from the new document."""

        old_doc = DocxDocument()
        old_doc.add_paragraph("Base text")
        old_path = self.tmpdir / "old_insert.docx"
        old_doc.save(old_path)

        new_doc = DocxDocument()
        paragraph = new_doc.add_paragraph("Base text")
        styled_run = paragraph.add_run(" styled addition")
        styled_run.font.bold = True
        styled_run.font.italic = True
        styled_run.font.underline = True
        styled_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        styled_run.font.name = "Arial"
        styled_run.font.size = Pt(18)
        styled_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        new_path = self.tmpdir / "new_insert.docx"
        new_doc.save(new_path)

        out_path = self.tmpdir / "out_insert.docx"
        make_redline_docx(str(old_path), str(new_path), str(out_path), author="Tester", date_iso="2025-01-03T00:00:00Z")

        doc_root = _read_part(out_path, "word/document.xml")
        inserted_runs = doc_root.xpath("//w:ins//w:r[w:t[contains(text(),'styled addition')]]", namespaces=NS)
        self.assertTrue(inserted_runs, "Inserted styled run should be present in output")
        ins_rpr = inserted_runs[0].find("w:rPr", namespaces=NS)
        self.assertIsNotNone(ins_rpr.find("w:b", namespaces=NS))
        self.assertIsNotNone(ins_rpr.find("w:i", namespaces=NS))
        self.assertIsNotNone(ins_rpr.find("w:u", namespaces=NS))
        self.assertIsNotNone(ins_rpr.find("w:highlight", namespaces=NS))
        ins_color = ins_rpr.find("w:color", namespaces=NS)
        self.assertEqual(ins_color.get(f"{{{NS['w']}}}val"), "0000FF")
        fonts = ins_rpr.find("w:rFonts", namespaces=NS)
        self.assertEqual(fonts.get(f"{{{NS['w']}}}ascii"), "Arial")
        size_el = ins_rpr.find("w:sz", namespaces=NS)
        self.assertEqual(size_el.get(f"{{{NS['w']}}}val"), "36")


if __name__ == "__main__":
    unittest.main()
