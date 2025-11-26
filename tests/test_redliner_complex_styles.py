#!/usr/bin/env python3
"""Complex styling tests for redliner with multiple styles per paragraph.

These tests focus on:
- Multiple runs with different formatting in the same paragraph
- Mixed styling changes within paragraphs
- Partial style changes affecting specific runs
- Style combinations and transformations
- Multi-paragraph scenarios with interconnected changes
"""

import os
import sys
import tempfile
import shutil
import zipfile
from pathlib import Path

# Add the parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from lxml import etree
from docx import Document as PythonDocx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from claude_office_skills.public.docx.scripts.document import Document
from claude_office_skills.public.docx.scripts.redliner import Redliner
from claude_office_skills.public.docx.ooxml.scripts.unpack import unpack_document
from claude_office_skills.public.docx.ooxml.scripts.pack import pack_document


NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': NS_W}


def qn(tag: str) -> str:
    """Expand a QName like 'w:p' to '{namespace}p'."""
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


class TestFixture:
    """Helper class for creating test documents."""

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def cleanup(self):
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def create_docx(self, name: str, creator_func) -> str:
        """Create a docx file using a creator function."""
        doc = PythonDocx()
        creator_func(doc)
        path = os.path.join(self.temp_dir, name)
        doc.save(path)
        return path

    def unpack_docx(self, docx_path: str) -> str:
        """Unpack a docx file and return the unpacked directory path."""
        unpacked_dir = docx_path.replace('.docx', '_unpacked')
        unpack_document(docx_path, unpacked_dir)
        return unpacked_dir

    def create_redlined_docx(self, old_path: str, new_path: str, author: str = 'TestAuthor') -> str:
        """Create a redlined docx by comparing old and new."""
        old_unpacked = self.unpack_docx(old_path)
        new_unpacked = self.unpack_docx(new_path)

        old_doc = Document(old_unpacked)
        new_doc = Document(new_unpacked)

        redliner = Redliner(old_doc, new_doc)
        redliner.redline(author=author)

        output_path = os.path.join(self.temp_dir, 'redlined.docx')
        new_doc.save(destination=new_unpacked, validate=False)
        pack_document(new_unpacked, output_path, validate=False)

        return output_path

    def parse_document_xml(self, docx_path: str) -> etree._Element:
        """Parse the document.xml from a docx file."""
        with zipfile.ZipFile(docx_path, 'r') as z:
            doc_xml = z.read('word/document.xml')
        return etree.fromstring(doc_xml)


@pytest.fixture
def test_fixture():
    """Pytest fixture for test setup and cleanup."""
    fixture = TestFixture()
    yield fixture
    fixture.cleanup()


# ============================================================================
# Multi-Run Paragraph Tests
# ============================================================================

class TestMultiRunParagraph:
    """Tests for paragraphs with multiple differently styled runs."""

    def test_three_runs_middle_format_change(self, test_fixture):
        """Test: Normal | Bold | Normal -> Normal | Italic | Normal"""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Start ')
            r = p.add_run('middle')
            r.bold = True
            p.add_run(' end')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Start ')
            r = p.add_run('middle')
            r.italic = True
            p.add_run(' end')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) >= 1, "Should track middle run format change"

    def test_five_runs_alternating_styles(self, test_fixture):
        """Test: Bold | Normal | Bold | Normal | Bold -> Normal | Bold | Normal | Bold | Normal"""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('One')
            r1.bold = True
            p.add_run(' Two ')
            r3 = p.add_run('Three')
            r3.bold = True
            p.add_run(' Four ')
            r5 = p.add_run('Five')
            r5.bold = True

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('One')
            r2 = p.add_run(' Two ')
            r2.bold = True
            p.add_run('Three')
            r4 = p.add_run(' Four ')
            r4.bold = True
            p.add_run('Five')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Should have format changes for all runs that changed
        assert len(rPrChanges) >= 3, "Should track multiple alternating format changes"

    def test_all_runs_same_change(self, test_fixture):
        """Test: All runs change from bold to italic."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('First ')
            r1.bold = True
            r2 = p.add_run('Second ')
            r2.bold = True
            r3 = p.add_run('Third')
            r3.bold = True

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('First ')
            r1.italic = True
            r2 = p.add_run('Second ')
            r2.italic = True
            r3 = p.add_run('Third')
            r3.italic = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) >= 3, "Should track format change for all three runs"


class TestMixedStylesWithTextChanges:
    """Tests combining text changes with style changes in multi-run paragraphs."""

    def test_text_change_in_styled_run(self, test_fixture):
        """Test text modification within a styled run."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Normal text ')
            r = p.add_run('bold original')
            r.bold = True
            p.add_run(' more normal')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Normal text ')
            r = p.add_run('bold modified')
            r.bold = True
            p.add_run(' more normal')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(insertions) > 0 or len(deletions) > 0, "Should track text change in bold run"

    def test_text_and_style_change_together(self, test_fixture):
        """Test both text and style changing in the same run."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Prefix ')
            r = p.add_run('old text')
            r.bold = True
            p.add_run(' suffix')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Prefix ')
            r = p.add_run('new text')
            r.italic = True
            p.add_run(' suffix')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        # Should have both insertions and deletions for text change
        total_changes = len(insertions) + len(deletions)
        assert total_changes > 0, "Should track text and style changes"

    def test_add_styled_run_between_existing(self, test_fixture):
        """Test inserting a new styled run between existing runs."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Beginning ')
            r = p.add_run('end')
            r.underline = True

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Beginning ')
            r1 = p.add_run('INSERTED ')
            r1.bold = True
            r1.italic = True
            r2 = p.add_run('end')
            r2.underline = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        assert len(insertions) > 0, "Should track inserted styled run"

    def test_remove_styled_run_from_middle(self, test_fixture):
        """Test removing a styled run from the middle."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Start ')
            r1 = p.add_run('REMOVE THIS ')
            r1.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.add_run('end')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('Start ')
            p.add_run('end')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(deletions) > 0, "Should track deleted styled run"


class TestComplexFormattingCombinations:
    """Tests for complex combinations of multiple formatting attributes."""

    def test_bold_italic_underline_to_color_size(self, test_fixture):
        """Test: Bold+Italic+Underline -> Red+Large"""
        def create_old(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')
            r.bold = True
            r.italic = True
            r.underline = True

        def create_new(doc):
            p = doc.add_paragraph()
            r = p.add_run('Styled text')
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r.font.size = Pt(24)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) > 0, "Should track complex format change"

        # Verify the previous state is recorded
        for change in rPrChanges:
            prior_rPr = change.find(qn('w:rPr'))
            assert prior_rPr is not None, "rPrChange should contain previous formatting"

    def test_multiple_runs_different_complex_changes(self, test_fixture):
        """Test multiple runs each with different complex formatting changes."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Bold ')
            r1.bold = True
            r2 = p.add_run('Italic ')
            r2.italic = True
            r3 = p.add_run('Underline ')
            r3.underline = True
            r4 = p.add_run('Red')
            r4.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Bold ')
            r1.italic = True  # Changed to italic
            r2 = p.add_run('Italic ')
            r2.bold = True  # Changed to bold
            r3 = p.add_run('Underline ')
            r3.font.strike = True  # Changed to strikethrough
            r4 = p.add_run('Red')
            r4.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Changed to blue

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # All 4 runs changed formatting
        assert len(rPrChanges) >= 4, "Should track format changes for all four runs"

    def test_gradient_of_sizes(self, test_fixture):
        """Test runs with gradually increasing sizes."""
        def create_old(doc):
            p = doc.add_paragraph()
            for size in [10, 12, 14, 16, 18]:
                r = p.add_run(f'Size{size} ')
                r.font.size = Pt(size)

        def create_new(doc):
            p = doc.add_paragraph()
            for size in [12, 14, 16, 18, 20]:  # All shifted up by 2pt
                r = p.add_run(f'Size{size-2} ')  # Keep same text
                r.font.size = Pt(size)

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) >= 5, "Should track size changes for all runs"


class TestPartialParagraphChanges:
    """Tests where only some runs in a paragraph change."""

    def test_first_run_unchanged_rest_changed(self, test_fixture):
        """Test first run stays same, rest change."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Unchanged ')
            r1.bold = True
            r2 = p.add_run('Changed1 ')
            r3 = p.add_run('Changed2')

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Unchanged ')
            r1.bold = True  # Same as old
            r2 = p.add_run('Changed1 ')
            r2.italic = True  # Added italic
            r3 = p.add_run('Changed2')
            r3.underline = True  # Added underline

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Only 2 runs should have changes
        assert len(rPrChanges) == 2, "Should track exactly 2 format changes"

    def test_last_run_unchanged_rest_changed(self, test_fixture):
        """Test last run stays same, rest change."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Changed1 ')
            r2 = p.add_run('Changed2 ')
            r3 = p.add_run('Unchanged')
            r3.bold = True

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Changed1 ')
            r1.italic = True  # Added italic
            r2 = p.add_run('Changed2 ')
            r2.underline = True  # Added underline
            r3 = p.add_run('Unchanged')
            r3.bold = True  # Same as old

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Only 2 runs should have changes
        assert len(rPrChanges) == 2, "Should track exactly 2 format changes"

    def test_middle_runs_unchanged(self, test_fixture):
        """Test edge runs change, middle stays same."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Edge1 ')
            r1.bold = True
            r2 = p.add_run('Middle ')
            r2.italic = True
            r3 = p.add_run('Edge2')
            r3.underline = True

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Edge1 ')
            # Changed from bold to normal
            r2 = p.add_run('Middle ')
            r2.italic = True  # Same
            r3 = p.add_run('Edge2')
            # Changed from underline to normal

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Only edge runs should have changes
        assert len(rPrChanges) == 2, "Should track exactly 2 format changes (edge runs)"


class TestMultiParagraphScenarios:
    """Tests with multiple paragraphs having different style changes."""

    def test_two_paragraphs_different_changes(self, test_fixture):
        """Test two paragraphs with completely different changes."""
        def create_old(doc):
            p1 = doc.add_paragraph()
            r1 = p1.add_run('Para1 ')
            r1.bold = True
            r1a = p1.add_run('bold text')
            r1a.bold = True

            p2 = doc.add_paragraph()
            r2 = p2.add_run('Para2 ')
            r2.italic = True
            r2a = p2.add_run('italic text')
            r2a.italic = True

        def create_new(doc):
            p1 = doc.add_paragraph()
            r1 = p1.add_run('Para1 ')
            r1.underline = True  # Changed from bold
            r1a = p1.add_run('bold text')
            r1a.underline = True  # Changed from bold

            p2 = doc.add_paragraph()
            r2 = p2.add_run('Para2 ')
            r2.font.strike = True  # Changed from italic
            r2a = p2.add_run('italic text')
            r2a.font.strike = True  # Changed from italic

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) >= 4, "Should track format changes in both paragraphs"

    def test_paragraph_added_with_styles(self, test_fixture):
        """Test adding a new paragraph with multiple styled runs."""
        def create_old(doc):
            p1 = doc.add_paragraph()
            r = p1.add_run('Existing paragraph')
            r.bold = True

        def create_new(doc):
            p1 = doc.add_paragraph()
            r = p1.add_run('Existing paragraph')
            r.bold = True

            p2 = doc.add_paragraph()
            r1 = p2.add_run('New ')
            r1.italic = True
            r2 = p2.add_run('styled ')
            r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r3 = p2.add_run('paragraph')
            r3.underline = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        insertions = doc.xpath('.//w:ins', namespaces=NSMAP)
        assert len(insertions) > 0, "Should track new paragraph as insertion"

    def test_paragraph_deleted_with_styles(self, test_fixture):
        """Test deleting a paragraph with multiple styled runs."""
        def create_old(doc):
            p1 = doc.add_paragraph()
            r = p1.add_run('Keep this')
            r.bold = True

            p2 = doc.add_paragraph()
            r1 = p2.add_run('Delete ')
            r1.italic = True
            r2 = p2.add_run('this ')
            r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r3 = p2.add_run('paragraph')
            r3.underline = True

        def create_new(doc):
            p1 = doc.add_paragraph()
            r = p1.add_run('Keep this')
            r.bold = True

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        deletions = doc.xpath('.//w:del', namespaces=NSMAP)
        assert len(deletions) > 0, "Should track deleted paragraph"


class TestStyleSwapping:
    """Tests for swapping styles between runs."""

    def test_swap_bold_italic_between_runs(self, test_fixture):
        """Test swapping bold and italic between two runs."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('FirstRun ')
            r1.bold = True
            r2 = p.add_run('SecondRun')
            r2.italic = True

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('FirstRun ')
            r1.italic = True  # Was bold
            r2 = p.add_run('SecondRun')
            r2.bold = True  # Was italic

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) == 2, "Should track both style swaps"

    def test_swap_colors_between_runs(self, test_fixture):
        """Test swapping colors between two runs."""
        def create_old(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Red ')
            r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r2 = p.add_run('Blue')
            r2.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Red ')
            r1.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Now blue
            r2 = p.add_run('Blue')
            r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Now red

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) == 2, "Should track both color swaps"


class TestRealWorldScenarios:
    """Tests based on real-world document editing scenarios."""

    def test_emphasis_added_to_key_terms(self, test_fixture):
        """Test adding bold to key terms in a sentence.

        When a single run is split into multiple runs with different styles,
        the formatting changes should be tracked with rPrChange.
        """
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('The contract includes payment terms, delivery conditions, and liability clauses.')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('The ')
            r1 = p.add_run('contract')
            r1.bold = True
            p.add_run(' includes ')
            r2 = p.add_run('payment terms')
            r2.bold = True
            p.add_run(', ')
            r3 = p.add_run('delivery conditions')
            r3.bold = True
            p.add_run(', and ')
            r4 = p.add_run('liability clauses')
            r4.bold = True
            p.add_run('.')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Should have rPrChange for each term that became bold (4 terms)
        assert len(rPrChanges) >= 4, "Should track bold added to key terms"

    def test_citation_styling(self, test_fixture):
        """Test citation formatting changes (italic book title, quotes)."""
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('As stated in ')
            r = p.add_run('The Art of War')
            p.add_run(' by Sun Tzu.')

        def create_new(doc):
            p = doc.add_paragraph()
            p.add_run('As stated in ')
            r = p.add_run('The Art of War')
            r.italic = True
            p.add_run(' by Sun Tzu.')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        assert len(rPrChanges) >= 1, "Should track italic added to book title"

    def test_warning_text_highlighting(self, test_fixture):
        """Test adding warning styling to important text.

        When run boundaries change and formatting is applied,
        the formatting changes should be tracked with rPrChange.
        """
        def create_old(doc):
            p = doc.add_paragraph()
            p.add_run('Important: Do not proceed without authorization. ')
            p.add_run('Contact your supervisor for approval.')

        def create_new(doc):
            p = doc.add_paragraph()
            r1 = p.add_run('Important: ')
            r1.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            r2 = p.add_run('Do not proceed without authorization. ')
            r2.bold = True
            p.add_run('Contact your supervisor for approval.')

        old_path = test_fixture.create_docx('old.docx', create_old)
        new_path = test_fixture.create_docx('new.docx', create_new)
        redlined_path = test_fixture.create_redlined_docx(old_path, new_path)

        doc = test_fixture.parse_document_xml(redlined_path)
        rPrChanges = doc.xpath('.//w:rPrChange', namespaces=NSMAP)
        # Should have rPrChange for the formatted runs
        assert len(rPrChanges) >= 2, "Should track warning styling changes"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
