import os
import subprocess
import sys
import tempfile
import unittest
import zipfile

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, RGBColor
from lxml import etree

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NSMAP = {'w': NS_W, 'r': NS_R}


def qn(tag: str) -> str:
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx_qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(docx_qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def parse_document_xml(docx_path: str) -> etree._Element:
    with zipfile.ZipFile(docx_path) as zf:
        xml_bytes = zf.read('word/document.xml')
    return etree.fromstring(xml_bytes)


class StylePreservationRedlineTest(unittest.TestCase):
    """Validate that redline generation preserves run styling and hyperlink wrappers."""

    def setUp(self):
        self.script = os.path.join('attempts', 'attempt_4.py')
        self.python = sys.executable
        self.tmp = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.tmp.cleanup()

    def run_redline(self, old_path: str, new_path: str) -> str:
        out_path = os.path.join(self.tmp.name, 'out.docx')
        result = subprocess.run(
            [self.python, self.script, old_path, new_path, out_path], capture_output=True, text=True
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return out_path

    def test_insert_preserves_combined_styles(self):
        old_doc = Document()
        old_doc.add_paragraph('Base ')
        old_path = os.path.join(self.tmp.name, 'old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        para = new_doc.add_paragraph('Base ')
        run = para.add_run('Styled')
        run.bold = True
        run.italic = True
        run.underline = WD_UNDERLINE.DOUBLE
        run.font.color.rgb = RGBColor(0xFF, 0x33, 0x00)
        run.font.size = Pt(16)
        run.font.name = 'Arial'
        new_path = os.path.join(self.tmp.name, 'new.docx')
        new_doc.save(new_path)

        redlined_path = self.run_redline(old_path, new_path)
        root = parse_document_xml(redlined_path)
        inserted_run = root.xpath('//w:ins//w:r[w:t="Styled"]', namespaces=NSMAP)
        self.assertTrue(inserted_run, 'Inserted styled run missing from output')
        rPr = inserted_run[0].find(qn('w:rPr'))
        self.assertIsNotNone(rPr)
        self.assertIsNotNone(rPr.find(qn('w:b')))
        self.assertIsNotNone(rPr.find(qn('w:i')))
        self.assertIsNotNone(rPr.find(qn('w:u')))
        self.assertIsNotNone(rPr.find(qn('w:color')))
        self.assertEqual(rPr.find(qn('w:sz')).get(qn('w:val')), '32')
        rFonts = rPr.find(qn('w:rFonts'))
        self.assertIsNotNone(rFonts)
        self.assertEqual(rFonts.get(qn('w:ascii')), 'Arial')

    def test_deletion_keeps_original_style(self):
        old_doc = Document()
        para = old_doc.add_paragraph()
        run = para.add_run('Gone')
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        old_path = os.path.join(self.tmp.name, 'old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_doc.add_paragraph('')
        new_path = os.path.join(self.tmp.name, 'new.docx')
        new_doc.save(new_path)

        redlined_path = self.run_redline(old_path, new_path)
        root = parse_document_xml(redlined_path)
        deleted_run = root.xpath('//w:del//w:r[w:delText="Gone"]', namespaces=NSMAP)
        self.assertTrue(deleted_run, 'Deleted run missing from output')
        rPr = deleted_run[0].find(qn('w:rPr'))
        self.assertIsNotNone(rPr)
        self.assertIsNotNone(rPr.find(qn('w:b')))
        color = rPr.find(qn('w:color'))
        self.assertIsNotNone(color)
        self.assertEqual(color.get(qn('w:val')), '0066CC')

    def test_hyperlink_wrapper_and_style_preserved(self):
        old_doc = Document()
        old_para = old_doc.add_paragraph('See ')
        add_hyperlink(old_para, 'link', 'http://example.com')
        old_path = os.path.join(self.tmp.name, 'old.docx')
        old_doc.save(old_path)

        new_doc = Document()
        new_para = new_doc.add_paragraph('See ')
        add_hyperlink(new_para, 'link updated', 'http://example.com')
        new_path = os.path.join(self.tmp.name, 'new.docx')
        new_doc.save(new_path)

        redlined_path = self.run_redline(old_path, new_path)
        root = parse_document_xml(redlined_path)
        hyperlink_nodes = root.xpath('//w:ins//w:hyperlink', namespaces=NSMAP)
        self.assertTrue(hyperlink_nodes, 'Inserted hyperlink wrapper missing')
        hyperlink = hyperlink_nodes[0]
        self.assertTrue(hyperlink.get(qn('r:id')))
        hyperlink_run = hyperlink.xpath('.//w:r', namespaces=NSMAP)[0]
        rPr = hyperlink_run.find(qn('w:rPr'))
        self.assertIsNotNone(rPr)
        rStyle = rPr.find(qn('w:rStyle'))
        self.assertIsNotNone(rStyle)
        self.assertEqual(rStyle.get(qn('w:val')), 'Hyperlink')


if __name__ == '__main__':
    unittest.main()
